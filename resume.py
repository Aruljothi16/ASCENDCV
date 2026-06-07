import os
import sqlite3
import bcrypt
import json
import re
from datetime import datetime, timedelta
import streamlit as st
import PyPDF2
import plotly.express as px
import plotly.graph_objects as go
import pandas as pd
import numpy as np
import time
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import io
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from docx import Document as DocxDocument
from docx.shared import Inches
DB_PATH = os.path.join(os.path.expanduser("~"), "ascendcv.db")

# Download required NLTK data
import nltk
nltk.download('punkt', quiet=True)
nltk.download('punkt_tab', quiet=True)
nltk.download('stopwords', quiet=True)
nltk.download('averaged_perceptron_tagger', quiet=True)

# Set page config
st.set_page_config(
    page_title="AscendCV - AI Resume Optimizer", 
    layout="wide", 
    page_icon="🚀",
    initial_sidebar_state="expanded"
)

# Enhanced Professional Purple Theme CSS with Modern Design
st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');
    
    :root {
        --primary: #7E57C2;
        --primary-dark: #5E35B1;
        --primary-light: #B39DDB;
        --secondary: #26A69A;
        --success: #66BB6A;
        --warning: #FFA726;
        --error: #EF5350;
        --bg: #FAFAFA;
        --card-bg: #FFFFFF;
        --text-primary: #2E2E3A;
        --text-secondary: #666676;
        --border: #E1E1E8;
        --gradient-primary: linear-gradient(135deg, #7E57C2 0%, #5E35B1 100%);
        --gradient-secondary: linear-gradient(135deg, #26A69A 0%, #00796B 100%);
        --gradient-success: linear-gradient(135deg, #66BB6A 0%, #4CAF50 100%);
        --gradient-warning: linear-gradient(135deg, #FFA726 0%, #F57C00 100%);
        --shadow-sm: 0 4px 6px -1px rgba(126, 87, 194, 0.1), 0 2px 4px -1px rgba(126, 87, 194, 0.06);
        --shadow-md: 0 10px 15px -3px rgba(126, 87, 194, 0.1), 0 4px 6px -2px rgba(126, 87, 194, 0.05);
        --shadow-lg: 0 20px 25px -5px rgba(126, 87, 194, 0.1), 0 10px 10px -5px rgba(126, 87, 194, 0.04);
        --shadow-xl: 0 25px 50px -12px rgba(126, 87, 194, 0.25);
    }
    
    * { margin: 0; padding: 0; box-sizing: border-box; }
    
    html, body, [class*="css"] {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
        background: linear-gradient(135deg, #FAFAFA 0%, #F5F5F5 100%);
        color: var(--text-primary);
        line-height: 1.6;
    }
    
    .main { 
        background: transparent; 
        padding: 0 !important; 
    }
    
    /* Enhanced Hero Section with Glass Morphism - CENTERED */
    .hero-section {
        background: var(--gradient-primary);
        padding: 4rem 2rem 3rem;
        border-radius: 0 0 3rem 3rem;
        margin-bottom: 3rem;
        box-shadow: var(--shadow-xl);
        position: relative;
        overflow: hidden;
        text-align: center;
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
    }
    
    .hero-section::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: 
            radial-gradient(circle at 20% 80%, rgba(255,255,255,0.1) 0%, transparent 50%),
            radial-gradient(circle at 80% 20%, rgba(255,255,255,0.05) 0%, transparent 50%),
            radial-gradient(circle at 40% 40%, rgba(255,255,255,0.08) 0%, transparent 50%);
        backdrop-filter: blur(40px);
    }
    
    .hero-content {
        position: relative;
        z-index: 2;
        text-align: center;
        color: white;
        max-width: 1200px;
        margin: 0 auto;
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
    }
    
    .hero-title {
        font-size: 3rem;
        font-weight: 900;
        margin-bottom: 1rem;
        background: linear-gradient(135deg, #FFFFFF, #E1E1E8);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-shadow: 0 4px 20px rgba(0,0,0,0.1);
        line-height: 1.1;
        text-align: center;
    }
    
    .hero-subtitle {
        font-size: 1.2rem;
        font-weight: 400;
        opacity: 0.95;
        margin-bottom: 2rem;
        max-width: 700px;
        line-height: 1.6;
        text-align: center;
    }
    
    .hero-features {
        display: flex;
        justify-content: center;
        gap: 1.5rem;
        flex-wrap: wrap;
        margin-top: 2rem;
        text-align: center;
    }
    
    .feature-badge {
        background: rgba(255,255,255,0.15);
        backdrop-filter: blur(20px);
        border: 1px solid rgba(255,255,255,0.2);
        padding: 0.8rem 1.5rem;
        border-radius: 2rem;
        font-weight: 600;
        font-size: 0.9rem;
        transition: all 0.3s ease;
    }
    
    .feature-badge:hover {
        background: rgba(255,255,255,0.25);
        transform: translateY(-3px);
    }
    
    /* Enhanced Cards with Glass Morphism */
    .card {
        background: rgba(255, 255, 255, 0.85);
        backdrop-filter: blur(20px);
        border-radius: 2rem;
        padding: 2.5rem;
        box-shadow: var(--shadow-lg);
        border: 1px solid rgba(255, 255, 255, 0.3);
        margin-bottom: 2rem;
        position: relative;
        overflow: hidden;
        transition: all 0.3s ease;
    }
    
    .card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 4px;
        background: var(--gradient-primary);
        border-radius: 2rem 2rem 0 0;
    }
    
    .card:hover {
        transform: translateY(-5px);
        box-shadow: var(--shadow-xl);
        background: rgba(255, 255, 255, 0.95);
    }
    
    /* Enhanced Metric Cards */
    .metric-card {
        background: rgba(255, 255, 255, 0.9);
        backdrop-filter: blur(20px);
        border-radius: 1.5rem;
        padding: 2rem 1.5rem;
        text-align: center;
        box-shadow: var(--shadow-md);
        border: 1px solid rgba(255, 255, 255, 0.3);
        margin-bottom: 1.5rem;
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
    }
    
    .metric-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 3px;
        background: var(--gradient-primary);
    }
    
    .metric-card:hover {
        transform: translateY(-5px);
        box-shadow: var(--shadow-lg);
        background: rgba(255, 255, 255, 0.95);
    }
    
    .metric-value {
        font-size: 2.5rem;
        font-weight: 800;
        background: var(--gradient-primary);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        display: block;
        margin-bottom: 0.5rem;
        line-height: 1;
    }
    
    .metric-label {
        color: var(--text-secondary);
        font-size: 0.9rem;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    
    /* Enhanced Profile Summary Card */
    .profile-summary-card {
        background: linear-gradient(135deg, rgba(126, 87, 194, 0.08), rgba(94, 53, 177, 0.05));
        border-left: 6px solid var(--primary);
        padding: 2.5rem;
        border-radius: 2rem;
        margin: 2rem 0;
        box-shadow: var(--shadow-lg);
        border: 1px solid rgba(126, 87, 194, 0.2);
        transition: all 0.3s ease;
    }
    
    .profile-summary-card:hover {
        transform: translateY(-3px);
        box-shadow: var(--shadow-xl);
    }
    
    .strength-badge {
        display: inline-block;
        padding: 0.8rem 1.5rem;
        border-radius: 2rem;
        font-weight: 700;
        font-size: 1.1rem;
        margin: 1rem 0;
        box-shadow: var(--shadow-md);
    }
    
    .strength-excellent {
        background: linear-gradient(135deg, #66BB6A 0%, #4CAF50 100%);
        color: white;
    }
    
    .strength-good {
        background: linear-gradient(135deg, #26A69A 0%, #00796B 100%);
        color: white;
    }
    
    .strength-moderate {
        background: linear-gradient(135deg, #FFA726 0%, #F57C00 100%);
        color: white;
    }
    
    .strength-gap {
        background: linear-gradient(135deg, #EF5350 0%, #E53935 100%);
        color: white;
    }
    
    /* Enhanced Skill Tags */
    .skill-tag {
        display: inline-block;
        background: var(--gradient-primary);
        color: white;
        padding: 0.6rem 1.2rem;
        border-radius: 2rem;
        font-size: 0.85rem;
        font-weight: 600;
        margin: 0.3rem;
        transition: all 0.3s ease;
        box-shadow: var(--shadow-sm);
        border: 1px solid rgba(255,255,255,0.2);
    }
    
    .skill-tag:hover {
        transform: translateY(-2px) scale(1.05);
        box-shadow: var(--shadow-md);
    }
    
    .missing-skill-tag {
        display: inline-block;
        background: linear-gradient(135deg, #EF5350, #E53935);
        color: white;
        padding: 0.6rem 1.2rem;
        border-radius: 2rem;
        font-size: 0.85rem;
        font-weight: 600;
        margin: 0.3rem;
        transition: all 0.3s ease;
        box-shadow: var(--shadow-sm);
        border: 1px solid rgba(255,255,255,0.2);
    }
    
    .missing-skill-tag:hover {
        transform: translateY(-2px) scale(1.05);
        box-shadow: var(--shadow-md);
    }
    
    /* Enhanced Analysis Cards */
    .analysis-card {
        background: rgba(255, 255, 255, 0.85);
        backdrop-filter: blur(20px);
        border-radius: 1.5rem;
        padding: 2rem;
        box-shadow: var(--shadow-md);
        border: 1px solid rgba(255, 255, 255, 0.3);
        margin: 1.5rem 0;
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
    }
    
    .analysis-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 3px;
        background: var(--gradient-primary);
    }
    
    .analysis-card:hover {
        transform: translateY(-3px);
        box-shadow: var(--shadow-lg);
        background: rgba(255, 255, 255, 0.95);
    }
    
    .feature-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
        gap: 2rem;
        margin: 2rem 0;
    }
    
    .feature-card {
        background: white;
        border-radius: 1rem;
        padding: 2rem;
        text-align: center;
        box-shadow: var(--shadow-md);
        border: 1px solid var(--border);
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
    }
    
    .feature-card:hover {
        transform: translateY(-4px);
        box-shadow: var(--shadow-xl);
    }
    
    .feature-icon {
        width: 4rem;
        height: 4rem;
        border-radius: 1rem;
        display: flex;
        align-items: center;
        justify-content: center;
        margin: 0 auto 1.5rem auto;
        font-size: 1.8rem;
        background: var(--gradient-primary);
        color: white;
    }
    
    .feature-title {
        font-size: 1.25rem;
        font-weight: 700;
        color: var(--text-primary);
        margin-bottom: 1rem;
    }
    
    .feature-desc {
        color: var(--text-secondary);
        line-height: 1.6;
        font-size: 0.95rem;
    }
    
    /* Enhanced Buttons */
    .stButton > button {
        background: var(--gradient-primary);
        color: white;
        border: none;
        border-radius: 1.2rem;
        padding: 1rem 2.5rem;
        font-weight: 600;
        transition: all 0.3s ease;
        box-shadow: var(--shadow-md);
        font-size: 1rem;
        position: relative;
        overflow: hidden;
    }
    
    .stButton > button::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(255,255,255,0.3), transparent);
        transition: left 0.5s;
    }
    
    .stButton > button:hover::before {
        left: 100%;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: var(--shadow-lg);
        background: var(--gradient-primary);
        color: white;
    }
    
    /* Enhanced Tabs */
    .stTabs [data-baseweb="tab-list"] {
        gap: 1rem;
        background: transparent;
        padding: 0.5rem;
        border-bottom: 2px solid rgba(126, 87, 194, 0.1);
    }
    
    .stTabs [data-baseweb="tab"] {
        height: 3.5rem;
        background: rgba(255, 255, 255, 0.8);
        backdrop-filter: blur(20px);
        border-radius: 1.2rem;
        padding: 0 2rem;
        font-weight: 600;
        border: 2px solid rgba(126, 87, 194, 0.1);
        transition: all 0.3s ease;
        box-shadow: var(--shadow-sm);
        font-size: 0.95rem;
    }
    
    .stTabs [data-baseweb="tab"]:hover {
        transform: translateY(-2px);
        box-shadow: var(--shadow-md);
        background: rgba(255, 255, 255, 0.9);
    }
    
    .stTabs [aria-selected="true"] {
        background: var(--gradient-primary);
        color: white;
        border-color: var(--primary);
        box-shadow: var(--shadow-md);
        transform: translateY(-2px);
    }
    
    /* Enhanced Suggestion Cards */
    .suggestion-card {
        background: linear-gradient(135deg, rgba(255,248,225,0.9), rgba(255,243,224,0.9));
        backdrop-filter: blur(20px);
        border-left: 5px solid #FFA726;
        padding: 1.5rem;
        border-radius: 1.5rem;
        margin: 1.5rem 0;
        box-shadow: var(--shadow-md);
        border: 1px solid rgba(255,255,255,0.3);
        transition: all 0.3s ease;
    }
    
    .suggestion-card:hover {
        transform: translateX(5px);
        box-shadow: var(--shadow-lg);
    }
    
    /* Enhanced Resource Cards */
    .resource-card {
        background: rgba(255, 255, 255, 0.85);
        backdrop-filter: blur(20px);
        border-radius: 1.5rem;
        padding: 1.5rem;
        box-shadow: var(--shadow-md);
        border: 1px solid rgba(255, 255, 255, 0.3);
        margin: 1.5rem 0;
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
    }
    
    .resource-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 3px;
        background: var(--gradient-secondary);
    }
    
    .resource-card:hover {
        transform: translateY(-5px);
        box-shadow: var(--shadow-lg);
        background: rgba(255, 255, 255, 0.95);
    }
    
    /* Enhanced Detail Cards */
    .skill-detail-card {
        background: linear-gradient(135deg, rgba(227,242,253,0.9), rgba(225,245,254,0.9));
        backdrop-filter: blur(20px);
        border-radius: 2rem;
        padding: 2rem;
        margin: 1.5rem 0;
        border-left: 6px solid #2196F3;
        box-shadow: var(--shadow-md);
        border: 1px solid rgba(255,255,255,0.3);
        transition: all 0.3s ease;
    }
    
    .skill-detail-card:hover {
        transform: translateY(-3px);
        box-shadow: var(--shadow-lg);
    }
    
    .learning-path-card {
        background: linear-gradient(135deg, rgba(232,245,232,0.9), rgba(241,248,233,0.9));
        backdrop-filter: blur(20px);
        border-radius: 2rem;
        padding: 2rem;
        margin: 1.5rem 0;
        border-left: 6px solid #4CAF50;
        box-shadow: var(--shadow-md);
        border: 1px solid rgba(255,255,255,0.3);
        transition: all 0.3s ease;
    }
    
    .learning-path-card:hover {
        transform: translateY(-3px);
        box-shadow: var(--shadow-lg);
    }
    
    .project-idea-card {
        background: linear-gradient(135deg, rgba(255,243,224,0.9), rgba(255,248,225,0.9));
        backdrop-filter: blur(20px);
        border-radius: 1.2rem;
        padding: 1.5rem;
        margin: 1rem 0;
        border-left: 4px solid #FF9800;
        box-shadow: var(--shadow-sm);
        border: 1px solid rgba(255,255,255,0.3);
        transition: all 0.3s ease;
    }
    
    .project-idea-card:hover {
        transform: translateX(5px);
        box-shadow: var(--shadow-md);
    }
    
    /* Quick Action Cards */
    .quick-action-card {
        background: linear-gradient(135deg, rgba(243,229,245,0.9), rgba(237,231,246,0.9));
        backdrop-filter: blur(20px);
        border-radius: 1.5rem;
        padding: 2rem;
        margin: 1.2rem 0;
        border-left: 5px solid var(--primary);
        box-shadow: var(--shadow-md);
        border: 1px solid rgba(255,255,255,0.3);
        transition: all 0.3s ease;
    }
    
    .quick-action-card:hover {
        transform: translateY(-5px);
        box-shadow: var(--shadow-lg);
    }
    
    /* Custom Scrollbar */
    ::-webkit-scrollbar {
        width: 10px;
    }
    
    ::-webkit-scrollbar-track {
        background: rgba(241,241,241,0.8);
        backdrop-filter: blur(20px);
        border-radius: 5px;
    }
    
    ::-webkit-scrollbar-thumb {
        background: var(--gradient-primary);
        border-radius: 5px;
        border: 2px solid rgba(255,255,255,0.3);
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: var(--primary-dark);
    }
    
    /* Animations */
    @keyframes slideIn {
        from {
            opacity: 0;
            transform: translateY(-20px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    @keyframes fadeIn {
        from {
            opacity: 0;
        }
        to {
            opacity: 1;
        }
    }
    
    /* File Uploader Enhancement */
    .stFileUploader>div>div {
        background: rgba(255, 255, 255, 0.9);
        backdrop-filter: blur(20px);
        border: 2px dashed rgba(126, 87, 194, 0.3);
        border-radius: 1.2rem;
        transition: all 0.3s ease;
    }
    
    .stFileUploader>div>div:hover {
        border-color: var(--primary);
        background: rgba(255, 255, 255, 0.95);
    }
    </style>
    """,
    unsafe_allow_html=True,
)
class AdvancedResumeAnalyzer:
    def __init__(self):
        self.stop_words = set(stopwords.words('english'))
        
        # Load projects dataset
        self.projects_dataset = self.load_projects_dataset()
        
        # Enhanced skill databases with learning resources and hyperlinks
        self.technical_skills = {
            'Programming Languages': [
                'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust',
                'swift', 'kotlin', 'scala', 'r', 'matlab', 'perl', 'dart', 'objective-c', 'shell', 'bash',
                'powershell', 'html', 'css', 'sql'
            ],
            'Frameworks & Libraries': [
                'react', 'angular', 'vue', 'django', 'flask', 'spring', 'laravel', 'express', 'node.js',
                'react native', 'flutter', 'tensorflow', 'pytorch', 'pandas', 'numpy', 'scikit-learn',
                'bootstrap', 'jquery', 'asp.net', 'ruby on rails', 'keras', 'opencv'
            ],
            'Databases': [
                'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'oracle', 'sql server', 'sqlite',
                'cassandra', 'dynamodb', 'firebase', 'firestore', 'cosmos db', 'bigquery', 'snowflake'
            ],
            'Cloud & DevOps': [
                'aws', 'amazon web services', 'azure', 'google cloud', 'gcp', 'docker', 'kubernetes',
                'jenkins', 'terraform', 'ansible', 'git', 'github', 'gitlab', 'ci/cd', 'linux', 'unix',
                'nginx', 'apache', 'prometheus', 'grafana'
            ],
            'Data Science & AI': [
                'machine learning', 'deep learning', 'data analysis', 'data visualization', 'nlp',
                'natural language processing', 'computer vision', 'tableau', 'power bi', 'spark', 'hadoop',
                'data mining', 'statistical analysis', 'business intelligence'
            ],
            'Tools & Platforms': [
                'jira', 'confluence', 'slack', 'teams', 'microsoft office', 'excel', 'word', 'powerpoint',
                'visual studio', 'vs code', 'intellij', 'eclipse', 'postman', 'figma', 'sketch', 'photoshop'
            ]
        }
        
        self.soft_skills = [
            'leadership', 'communication', 'teamwork', 'collaboration', 'problem solving', 'critical thinking',
            'creativity', 'adaptability', 'time management', 'project management', 'analytical thinking',
            'decision making', 'strategic thinking', 'attention to detail', 'multitasking', 'organization',
            'presentation', 'negotiation', 'conflict resolution', 'emotional intelligence', 'mentoring',
            'coaching', 'public speaking', 'stakeholder management', 'agile', 'scrum', 'kanban'
        ]

        # ONE-WEEK INTENSIVE LEARNING PLANS WITH HYPERLINKS
        self.skill_improvement_guides = {
            'python': {
                'priority': 'High',
                'time_estimate': '1 week',
                'resources': [
                    '[Python Crash Course - freeCodeCamp (YouTube)](https://www.youtube.com/watch?v=rfscVS0vtbw)',
                    '[Python for Everybody - Coursera](https://www.coursera.org/specializations/python)',
                    '[Real Python Tutorials](https://realpython.com/)',
                    '[Python Official Documentation](https://docs.python.org/3/)',
                    '[Automate the Boring Stuff with Python](https://automatetheboringstuff.com/)'
                ],
                'practice_platforms': [
                    '[LeetCode Python](https://leetcode.com/problemset/all/)',
                    '[HackerRank Python](https://www.hackerrank.com/domains/python)',
                    '[Codewars Python](https://www.codewars.com/?language=python)',
                    '[Exercism Python](https://exercism.org/tracks/python)'
                ],
                'learning_path': [
                    'Day 1: Python syntax, variables, data types (6-8 hours) - Complete basic exercises',
                    'Day 2: Control flow, loops, functions (6-8 hours) - Build simple programs',
                    'Day 3: Data structures (lists, dicts, sets, tuples) (6-8 hours) - Practice data manipulation',
                    'Day 4: OOP concepts, classes, inheritance (6-8 hours) - Create class structures',
                    'Day 5: File handling, error handling, modules (6-8 hours) - Work with external files',
                    'Day 6: Build mini-project - Web scraper or data analyzer (8 hours)',
                    'Day 7: Complete full project - REST API or automation tool (8 hours)'
                ],
                'project_ideas': [
                    'Web scraper for news articles using BeautifulSoup',
                    'Data analysis with Pandas on COVID-19 dataset',
                    'Automation script for file management and organization',
                    'REST API with Flask for todo application',
                    'Weather app with API integration',
                    'Password manager with encryption'
                ],
                'quick_wins': [
                    'Complete Python Crash Course on YouTube (Day 1-2)',
                    'Solve 10 easy problems daily on HackerRank',
                    'Build one small automation script per day',
                    'Join Python Discord communities for quick help'
                ],
                'certifications': [
                    '[PCAP - Python Institute](https://pythoninstitute.org/pcap)',
                    '[Microsoft Python Certification](https://docs.microsoft.com/en-us/learn/certifications/)'
                ]
            },
            'javascript': {
                'priority': 'High',
                'time_estimate': '1 week',
                'resources': [
                    '[JavaScript Tutorial - JavaScript.info](https://javascript.info/)',
                    '[MDN JavaScript Guide](https://developer.mozilla.org/en-US/docs/Web/JavaScript)',
                    '[freeCodeCamp JavaScript](https://www.freecodecamp.org/learn/javascript-algorithms-and-data-structures/)',
                    '[Eloquent JavaScript](https://eloquentjavascript.net/)',
                    '[JavaScript30 Challenge](https://javascript30.com/)'
                ],
                'practice_platforms': [
                    '[LeetCode JavaScript](https://leetcode.com/problemset/all/)',
                    '[HackerRank JavaScript](https://www.hackerrank.com/domains/tutorials/10-days-of-javascript)',
                    '[Codewars JavaScript](https://www.codewars.com/?language=javascript)',
                    '[Frontend Mentor](https://www.frontendmentor.io/)'
                ],
                'learning_path': [
                    'Day 1: JS fundamentals, variables, data types (6-8 hours)',
                    'Day 2: Functions, scope, arrays, objects (6-8 hours)',
                    'Day 3: DOM manipulation, events, ES6+ features (6-8 hours)',
                    'Day 4: Async programming, promises, async/await (6-8 hours)',
                    'Day 5: API integration, fetch, JSON handling (6-8 hours)',
                    'Day 6: Build interactive web app (8 hours)',
                    'Day 7: Complete portfolio project with deployment (8 hours)'
                ],
                'project_ideas': [
                    'Interactive todo list with local storage',
                    'Weather app with geolocation API',
                    'Quiz application with timer and scoring',
                    'E-commerce product filter and cart',
                    'Real-time chat application',
                    'Music player with playlist'
                ],
                'quick_wins': [
                    'Complete JavaScript30 challenge - 30 projects in 30 days',
                    'Build one interactive component daily',
                    'Practice DOM manipulation daily',
                    'Join JavaScript Discord communities'
                ],
                'certifications': [
                    '[JavaScript Developer Certification](https://www.freecodecamp.org/learn/javascript-algorithms-and-data-structures/)'
                ]
            },
            'react': {
                'priority': 'High',
                'time_estimate': '1 week',
                'resources': [
                    '[React Official Tutorial](https://reactjs.org/tutorial/tutorial.html)',
                    '[React Documentation](https://reactjs.org/docs/getting-started.html)',
                    '[Scrimba React Course](https://scrimba.com/learn/learnreact)',
                    '[Epic React by Kent C. Dodds](https://epicreact.dev/)',
                    '[React Tutorial by freeCodeCamp](https://www.freecodecamp.org/news/tag/react/)'
                ],
                'practice_platforms': [
                    '[React Challenges](https://github.com/enaqx/awesome-react#react-tutorials)',
                    '[Frontend Mentor React](https://www.frontendmentor.io/challenges?difficulties=1&types=free&languages=react)',
                    '[Build React Components Daily](https://react.challenges.cc/)',
                    '[React Practice Projects](https://github.com/iamshaunjp/React-Firebase)'
                ],
                'learning_path': [
                    'Day 1: React basics, JSX, components (6-8 hours)',
                    'Day 2: State, props, event handling (6-8 hours)',
                    'Day 3: Hooks (useState, useEffect, useContext) (6-8 hours)',
                    'Day 4: Forms, conditional rendering, lists (6-8 hours)',
                    'Day 5: Routing, API integration, context (6-8 hours)',
                    'Day 6: Build complete React app (8 hours)',
                    'Day 7: Deploy, optimize, add features (8 hours)'
                ],
                'project_ideas': [
                    'Portfolio website with React and animations',
                    'Task management application with drag-drop',
                    'Movie database browser with search',
                    'Social media dashboard with charts',
                    'E-commerce product catalog',
                    'Blog with markdown support'
                ],
                'quick_wins': [
                    'Scrimba React course - interactive learning',
                    'Clone popular UI components daily',
                    'Build one feature per day',
                    'Deploy projects on Vercel/Netlify'
                ],
                'certifications': [
                    '[Meta React Developer Certificate](https://www.coursera.org/professional-certificates/meta-react-developer)'
                ]
            },
            'node.js': {
                'priority': 'High',
                'time_estimate': '1 week',
                'resources': [
                    '[Node.js Official Docs](https://nodejs.org/en/docs/)',
                    '[The Net Ninja Node.js](https://www.youtube.com/playlist?list=PL4cUxeGkcC9gcy9lrvMJ75z9maRw4byYp)',
                    '[Node.js Tutorial by freeCodeCamp](https://www.freecodecamp.org/news/tag/node/)',
                    '[Express.js Guide](https://expressjs.com/en/guide/routing.html)',
                    '[MongoDB with Node.js](https://www.mongodb.com/languages/express-mongodb-rest-api-tutorial)'
                ],
                'practice_platforms': [
                    '[Build REST APIs Daily](https://github.com/public-apis/public-apis)',
                    '[Node.js Best Practices](https://github.com/goldbergyoni/nodebestpractices)',
                    '[API Challenge Platform](https://www.postman.com/company/student-program/)',
                    '[Backend Practice Projects](https://github.com/topics/backend-project)'
                ],
                'learning_path': [
                    'Day 1: Node.js basics, npm, modules (6-8 hours)',
                    'Day 2: Express.js, routing, middleware (6-8 hours)',
                    'Day 3: MongoDB integration, Mongoose (6-8 hours)',
                    'Day 4: RESTful API development (6-8 hours)',
                    'Day 5: Authentication, JWT, security (6-8 hours)',
                    'Day 6: Build backend project (8 hours)',
                    'Day 7: Testing, deployment, documentation (8 hours)'
                ],
                'project_ideas': [
                    'RESTful API for blog with CRUD operations',
                    'Authentication system with JWT tokens',
                    'Real-time chat server with Socket.io',
                    'File upload service with cloud storage',
                    'E-commerce backend with payment integration',
                    'URL shortener service'
                ],
                'quick_wins': [
                    'Build simple CRUD API in first 2 days',
                    'Practice with Postman daily',
                    'Deploy to Heroku/Railway',
                    'Join Node.js communities'
                ],
                'certifications': [
                    '[Node.js Certified Developer](https://openjsf.org/certification/)'
                ]
            },
            'sql': {
                'priority': 'High',
                'time_estimate': '1 week',
                'resources': [
                    '[SQLZoo Interactive Tutorial](https://sqlzoo.net/)',
                    '[SQLBolt](https://sqlbolt.com/)',
                    '[Mode Analytics SQL](https://mode.com/sql-tutorial/)',
                    '[Khan Academy SQL](https://www.khanacademy.org/computing/computer-programming/sql)',
                    '[W3Schools SQL](https://www.w3schools.com/sql/)'
                ],
                'practice_platforms': [
                    '[LeetCode SQL](https://leetcode.com/problemset/database/)',
                    '[HackerRank SQL](https://www.hackerrank.com/domains/sql)',
                    '[StrataScratch](https://www.stratascratch.com/)',
                    '[SQL Practice Online](https://www.sql-practice.com/)'
                ],
                'learning_path': [
                    'Day 1: SQL basics, SELECT queries (6-8 hours)',
                    'Day 2: WHERE, ORDER BY, filtering (6-8 hours)',
                    'Day 3: JOINs and relationships (6-8 hours)',
                    'Day 4: Aggregate functions, GROUP BY (6-8 hours)',
                    'Day 5: Subqueries, CTEs, window functions (6-8 hours)',
                    'Day 6: Database design, normalization (8 hours)',
                    'Day 7: Practice 50+ complex problems (8 hours)'
                ],
                'project_ideas': [
                    'Library management database with complex queries',
                    'E-commerce analytics with sales reports',
                    'Employee management system with payroll',
                    'Inventory tracking with alerts',
                    'Social media analytics dashboard',
                    'Banking transaction analysis'
                ],
                'quick_wins': [
                    'Complete SQLZoo tutorials in 2 days',
                    'Solve 10 SQL problems daily',
                    'Practice with real datasets',
                    'Build database schemas daily'
                ],
                'certifications': [
                    '[Microsoft SQL Certification](https://docs.microsoft.com/en-us/learn/certifications/sql-server-2016/)',
                    '[Oracle SQL Certification](https://education.oracle.com/oracle-database-sql-certified-associate/trackp_457)'
                ]
            }
        }
        
        # Default improvement plan
        self.default_improvement_plan = {
            'priority': 'Medium',
            'time_estimate': '1 week',
            'learning_path': [
                'Day 1: Fundamentals and basics (6-8 hours)',
                'Day 2-3: Core concepts and practice (6-8 hours/day)',
                'Day 4-5: Advanced topics and projects (6-8 hours/day)',
                'Day 6-7: Build complete project (8 hours/day)'
            ],
            'project_ideas': [
                'Build practical application',
                'Create portfolio project',
                'Solve real-world problem',
                'Contribute to open source'
            ]
        }
        
        # Flatten all technical skills
        self.all_technical_skills = []
        for category_skills in self.technical_skills.values():
            self.all_technical_skills.extend(category_skills)

    def load_projects_dataset(self):
        """Load projects dataset from JSON file"""
        try:
            with open('projects_dataset.json', 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            return self.create_default_projects_dataset()

    def create_default_projects_dataset(self):
        """Create default projects dataset"""
        projects_dataset = {
            "web_development": [
                {
                    "name": "E-Commerce Platform",
                    "skills": ["html", "css", "javascript", "node.js", "mongodb", "express"],
                    "bullets": [
                        "Implemented user authentication system with JWT tokens, shopping cart functionality, and payment gateway integration using Stripe API",
                        "Designed and deployed RESTful API architecture with Express.js backend, supporting 1000+ concurrent users with 99.9% uptime",
                        "Increased user engagement by 45% through responsive design implementation and optimized database queries reducing load time by 60%"
                    ]
                },
                {
                    "name": "College Forum",
                    "skills": ["html", "css", "javascript", "php", "mysql"],
                    "bullets": [
                        "Implemented user registration/login, Q&A discussions, private messaging, and profile management",
                        "Designed and deployed a lightweight admin panel for monitoring discussions and moderating content",
                        "Increased student engagement and collaboration by 40% through centralized Q&A discussions and achievement sharing"
                    ]
                }
            ],
            "data_science": [
                {
                    "name": "Customer Sentiment Analysis System",
                    "skills": ["python", "nltk", "scikit-learn", "pandas", "flask"],
                    "bullets": [
                        "Developed NLP-based sentiment analysis system processing 10K+ customer reviews daily with 89% classification accuracy",
                        "Implemented text preprocessing pipeline with tokenization, lemmatization, and TF-IDF vectorization for feature extraction",
                        "Reduced customer complaint response time by 45% through automated sentiment detection and priority routing"
                    ]
                }
            ],
            "machine_learning": [
                {
                    "name": "Image Classification System",
                    "skills": ["python", "tensorflow", "keras", "opencv", "flask"],
                    "bullets": [
                        "Developed CNN-based image classification model using transfer learning (ResNet50) achieving 94% accuracy on 50K+ images",
                        "Implemented data augmentation techniques and regularization methods reducing overfitting and improving model generalization",
                        "Deployed model as REST API using Flask handling 500+ requests/second with average response time under 200ms"
                    ]
                }
            ]
        }
        
        # Save to file for future use
        try:
            with open('projects_dataset.json', 'w', encoding='utf-8') as f:
                json.dump(projects_dataset, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"Warning: Could not save projects dataset: {e}")
        
        return projects_dataset

    def generate_projects_for_skills(self, matching_skills, missing_skills, num_projects=3):
        """Generate project suggestions based on matching and missing skills"""
        all_skills = list(set(matching_skills + missing_skills))
        generated_projects = []
        
        # Search through projects dataset
        for category, projects in self.projects_dataset.items():
            for project in projects:
                project_skills = [s.lower() for s in project['skills']]
                
                # Check how many skills match
                matching_count = sum(1 for s in all_skills if s.lower() in project_skills)
                
                if matching_count >= 2:  # At least 2 skills match
                    # Prioritize projects with both matching and missing skills
                    has_missing = any(s.lower() in project_skills for s in missing_skills)
                    has_matching = any(s.lower() in project_skills for s in matching_skills)
                    
                    if has_missing and has_matching:
                        generated_projects.append({
                            'name': project['name'],
                            'skills': project['skills'],
                            'bullets': project['bullets'],
                            'priority': 'high',
                            'match_score': matching_count
                        })
                    elif has_missing:
                        generated_projects.append({
                            'name': project['name'],
                            'skills': project['skills'],
                            'bullets': project['bullets'],
                            'priority': 'medium',
                            'match_score': matching_count
                        })
        
        # Sort by priority and match score
        generated_projects.sort(key=lambda x: (x['priority'] == 'high', x['match_score']), reverse=True)
        
        return generated_projects[:num_projects]

    def preprocess_text(self, text):
        """Clean and preprocess text"""
        if not text:
            return ""
        text = re.sub(r'\s+', ' ', text.lower().strip())
        text = re.sub(r'[^\w\s\-\.]', ' ', text)
        return text

    def extract_skills_with_context(self, text):
        """Extract skills with context awareness"""
        text = self.preprocess_text(text)
        found_skills = {
            'technical': [],
            'soft': [],
            'by_category': {}
        }
        
        # Extract technical skills
        for category, skills in self.technical_skills.items():
            category_found = []
            for skill in skills:
                patterns = [
                    r'\b' + re.escape(skill) + r'\b',
                    r'\b' + re.escape(skill.replace(' ', r'\s+')) + r'\b',
                    r'\b' + re.escape(skill.replace('.', r'\.?')) + r'\b'
                ]
                
                for pattern in patterns:
                    if re.search(pattern, text):
                        if skill not in category_found:
                            category_found.append(skill)
                        if skill not in found_skills['technical']:
                            found_skills['technical'].append(skill)
                        break
            
            if category_found:
                found_skills['by_category'][category] = category_found
        
        # Extract soft skills
        for skill in self.soft_skills:
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text) and skill not in found_skills['soft']:
                found_skills['soft'].append(skill)
        
        return found_skills

    def extract_experience_years(self, text):
        """Enhanced experience extraction"""
        patterns = [
            r'(\d+)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?(?:professional|relevant|work)?\s*experience',
            r'experience\s+of\s+(\d+)\+?\s*(?:years?|yrs?)',
            r'(\d+)\+?\s*(?:years?|yrs?)\s+in\s+(?:the\s+)?\w+',
            r'(\d+)\+?\s*(?:years?|yrs?)\s+(?:working|developing|building|managing)',
        ]
        
        years = []
        for pattern in patterns:
            matches = re.findall(pattern, text.lower())
            for match in matches:
                if match.isdigit():
                    years.append(int(match))
        
        # Look for date ranges
        date_pattern = r'(20\d{2})\s*[-–—]\s*(20\d{2}|present|current)'
        date_matches = re.findall(date_pattern, text)
        for start, end in date_matches:
            end_year = datetime.now().year if end in ['present', 'current'] else int(end)
            experience = end_year - int(start)
            if experience > 0:
                years.append(experience)
        
        return max(years) if years else 0

    def extract_education_level(self, text):
        """Extract education level"""
        text_lower = text.lower()
        
        education_keywords = {
            'phd': ['phd', 'ph.d', 'doctorate', 'doctoral'],
            'masters': ['master', 'masters', "master's", 'msc', 'm.sc', 'ma', 'm.a', 'mba'],
            'bachelors': ['bachelor', 'bachelors', "bachelor's", 'bsc', 'b.sc', 'ba', 'b.a', 'b.tech'],
            'associates': ['associate', 'associates', "associate's", 'diploma']
        }
        
        for level, keywords in education_keywords.items():
            for keyword in keywords:
                if re.search(r'\b' + re.escape(keyword) + r'\b', text_lower):
                    return level
        
        return 'not_specified'

    def calculate_advanced_keyword_match(self, resume_text, job_text):
        """Advanced keyword matching with TF-IDF"""
        if not resume_text or not job_text:
            return 0
        
        resume_clean = self.preprocess_text(resume_text)
        job_clean = self.preprocess_text(job_text)
        
        vectorizer = TfidfVectorizer(
            max_features=2000,
            stop_words='english',
            ngram_range=(1, 3),
            min_df=1,
            max_df=0.85,
            sublinear_tf=True
        )
        
        try:
            tfidf_matrix = vectorizer.fit_transform([resume_clean, job_clean])
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            return min(similarity * 100, 100)
        except:
            return 0

    def extract_achievements(self, text):
        """Extract quantified achievements"""
        patterns = [
            r'increased?\s+(?:\w+\s+){0,3}by\s+(\d+%|\$\d+[kmb]?|\d+(?:\.\d+)?[kmb]?)',
            r'reduced?\s+(?:\w+\s+){0,3}by\s+(\d+%|\$\d+[kmb]?|\d+(?:\.\d+)?[kmb]?)',
            r'improved?\s+(?:\w+\s+){0,3}by\s+(\d+%|\d+(?:\.\d+)?[kmb]?)',
            r'saved?\s+(?:\w+\s+){0,3}(?:of\s+)?(\$\d+[kmb]?|\d+[kmb]?)',
            r'generated?\s+(?:\w+\s+){0,3}(?:of\s+)?(\$\d+[kmb]?|\d+[kmb]?)',
            r'managed?\s+(?:a\s+)?(?:\w+\s+){0,2}(?:of\s+)?(\$\d+[kmb]?|\d+[kmb]?)',
        ]
        
        achievements = []
        sentences = sent_tokenize(text)
        
        for sentence in sentences:
            for pattern in patterns:
                if re.search(pattern, sentence.lower()):
                    achievements.append(sentence.strip())
                    break
        
        return list(set(achievements))[:10]

    def generate_profile_summary(self, analysis, resume_text):
        """Generate comprehensive profile summary with enhanced insights"""
        summary = {
            'strength_level': '',
            'strength_class': '',
            'key_strengths': [],
            'areas_for_improvement': [],
            'competitive_advantage': [],
            'recommendations': [],
            'readiness_score': 0,
            'profile_highlights': {
                'technical_coverage': 0,
                'experience_level': '',
                'achievement_quality': '',
                'overall_readiness': ''
            }
        }
        
        # Determine strength level and class
        score = analysis['match_percentage']
        if score >= 80:
            summary['strength_level'] = "⭐ Excellent Match"
            summary['strength_class'] = "strength-excellent"
            summary['competitive_advantage'].append("Your profile is highly competitive for this role")
            summary['readiness_score'] = 95
        elif score >= 65:
            summary['strength_level'] = "✅ Good Match"
            summary['strength_class'] = "strength-good"
            summary['competitive_advantage'].append("You have a strong foundation for this role")
            summary['readiness_score'] = 75
        elif score >= 50:
            summary['strength_level'] = "⚠️ Moderate Match"
            summary['strength_class'] = "strength-moderate"
            summary['competitive_advantage'].append("Focus on skill development to become competitive")
            summary['readiness_score'] = 55
        else:
            summary['strength_level'] = "📊 Skills Gap"
            summary['strength_class'] = "strength-gap"
            summary['competitive_advantage'].append("Significant upskilling needed for this role")
            summary['readiness_score'] = 30
        
        # Calculate technical coverage
        tech_coverage = (analysis['technical_skills_matched'] / max(analysis['technical_skills_total'], 1)) * 100
        summary['profile_highlights']['technical_coverage'] = round(tech_coverage, 1)
        
        # Experience level assessment
        exp_years = analysis['experience_years']
        if exp_years >= 5:
            summary['profile_highlights']['experience_level'] = "Senior Level"
        elif exp_years >= 3:
            summary['profile_highlights']['experience_level'] = "Mid Level"
        elif exp_years >= 1:
            summary['profile_highlights']['experience_level'] = "Junior Level"
        else:
            summary['profile_highlights']['experience_level'] = "Entry Level"
        
        # Achievement quality
        achievements = analysis['achievements_count']
        if achievements >= 5:
            summary['profile_highlights']['achievement_quality'] = "Outstanding"
        elif achievements >= 3:
            summary['profile_highlights']['achievement_quality'] = "Good"
        elif achievements >= 1:
            summary['profile_highlights']['achievement_quality'] = "Adequate"
        else:
            summary['profile_highlights']['achievement_quality'] = "Needs Improvement"
        
        # Overall readiness
        if summary['readiness_score'] >= 80:
            summary['profile_highlights']['overall_readiness'] = "Interview Ready"
        elif summary['readiness_score'] >= 60:
            summary['profile_highlights']['overall_readiness'] = "Nearly Ready"
        else:
            summary['profile_highlights']['overall_readiness'] = "Needs Development"
        
        # Key strengths
        if analysis['technical_skills_matched'] > 0:
            summary['key_strengths'].append(f"✅ {analysis['technical_skills_matched']} matching technical skills ({tech_coverage:.0f}% coverage)")
        if analysis['experience_years'] > 0:
            summary['key_strengths'].append(f"✅ {analysis['experience_years']} years of relevant experience")
        if analysis['achievements_count'] > 0:
            summary['key_strengths'].append(f"✅ {analysis['achievements_count']} quantified achievements demonstrating impact")
        if analysis['soft_skills_matched'] > 0:
            summary['key_strengths'].append(f"✅ {analysis['soft_skills_matched']} matching soft skills")
        if analysis['keyword_score'] >= 60:
            summary['key_strengths'].append(f"✅ Strong keyword alignment ({analysis['keyword_score']:.1f}%)")
        
        # Areas for improvement
        missing_count = len(analysis['missing_technical_skills'])
        if missing_count > 0:
            priority_count = sum(1 for s in analysis['missing_technical_skills'] if analysis['improvement_plan'].get(s, {}).get('priority') == 'High')
            summary['areas_for_improvement'].append(f"⚡ {missing_count} technical skills to acquire ({priority_count} high priority)")
        
        if tech_coverage < 70:
            summary['areas_for_improvement'].append(f"⚡ Technical skill coverage at {tech_coverage:.0f}% - aim for 80%+")
        
        if analysis['achievements_count'] < 3:
            summary['areas_for_improvement'].append("⚡ Add more quantified achievements to stand out (target: 5+)")
        
        if analysis['keyword_score'] < 60:
            summary['areas_for_improvement'].append(f"⚡ Keyword alignment needs improvement ({analysis['keyword_score']:.1f}%)")
        
        # Personalized recommendations with priority
        if missing_count > 0:
            priority_skills = [s for s in analysis['missing_technical_skills'] if 
                             analysis['improvement_plan'].get(s, {}).get('priority') == 'High']
            if priority_skills:
                summary['recommendations'].append(f"🎯 **High Priority:** Master these skills first - {', '.join(priority_skills[:3])}")
        
        if score < 75:
            summary['recommendations'].append("📚 **Action:** Follow the intensive learning schedule in the Learning Path tab")
        
        if analysis['experience_years'] < 2:
            summary['recommendations'].append("💼 **Portfolio:** Build 3-5 substantial projects to demonstrate practical experience")
        
        if analysis['achievements_count'] < 3:
            summary['recommendations'].append("📊 **Impact:** Quantify your achievements with metrics (percentages, numbers, dollar amounts)")
        
        if analysis['keyword_score'] < 70:
            summary['recommendations'].append("🔍 **Optimization:** Incorporate more keywords from the job description naturally")
        
        return summary

    def get_skill_improvement_plan(self, missing_skills):
        """Generate personalized improvement plan"""
        improvement_plan = {}
        
        for skill in missing_skills:
            if skill in self.skill_improvement_guides:
                improvement_plan[skill] = self.skill_improvement_guides[skill]
            else:
                improvement_plan[skill] = self.default_improvement_plan.copy()
                improvement_plan[skill]['resources'] = [
                    f'[Online courses for {skill.title()} on Coursera](https://www.coursera.org/search?query={skill})',
                    f'[Udemy {skill.title()} courses](https://www.udemy.com/courses/search/?q={skill})',
                    f'[edX {skill.title()} courses](https://www.edx.org/search?q={skill})',
                    f'Official {skill.title()} documentation'
                ]
        
        return improvement_plan

    def analyze_resume_advanced(self, resume_text, job_description):
        """Comprehensive resume analysis with enhanced accuracy"""
        if not resume_text or not job_description:
            return self._create_empty_analysis()
        
        try:
            # Extract skills
            resume_skills = self.extract_skills_with_context(resume_text)
            job_skills = self.extract_skills_with_context(job_description)
            
            # Calculate matches
            technical_resume = set(resume_skills['technical'])
            technical_job = set(job_skills['technical'])
            technical_match = len(technical_resume.intersection(technical_job))
            technical_total = len(technical_job) if technical_job else 1
            
            soft_resume = set(resume_skills['soft'])
            soft_job = set(job_skills['soft'])
            soft_match = len(soft_resume.intersection(soft_job))
            soft_total = len(soft_job) if soft_job else 1
            
            # Keyword matching
            keyword_score = self.calculate_advanced_keyword_match(resume_text, job_description)
            
            # Extract metrics
            experience_years = self.extract_experience_years(resume_text)
            education_level = self.extract_education_level(resume_text)
            achievements = self.extract_achievements(resume_text)
            
            # Generate projects based on skills
            generated_projects = self.generate_projects_for_skills(
                list(technical_resume.intersection(technical_job)),
                list(technical_job - technical_resume),
                num_projects=3
            )
            
            # Enhanced ATS score calculation with refined weights
            technical_weight = 0.40  # Increased importance
            keyword_weight = 0.25
            soft_skill_weight = 0.15
            experience_weight = 0.10
            education_weight = 0.05
            achievement_weight = 0.05
            
            # Technical score with bonus for exceeding requirements
            tech_match_rate = technical_match / max(technical_total, 1)
            technical_score = min(tech_match_rate * 120, 100) * technical_weight  # 120 allows for bonus
            
            # Keyword score
            keyword_score_weighted = keyword_score * keyword_weight
            
            # Soft skills score
            soft_match_rate = soft_match / max(soft_total, 1)
            soft_score = min(soft_match_rate * 110, 100) * soft_skill_weight
            
            # Experience score with diminishing returns
            exp_score_map = {0: 0, 1: 20, 2: 40, 3: 60, 4: 75, 5: 85, 6: 90, 7: 95, 8: 98}
            exp_base_score = exp_score_map.get(min(experience_years, 8), 100)
            experience_score = exp_base_score * experience_weight
            
            # Education score
            edu_score_map = {'phd': 100, 'masters': 85, 'bachelors': 70, 'associates': 50, 'not_specified': 30}
            education_score = edu_score_map.get(education_level, 30) * education_weight
            
            # Achievement bonus
            achievement_bonus = min(len(achievements) * 15, 100) * achievement_weight
            
            # Calculate total score
            total_score = (
                technical_score + 
                keyword_score_weighted + 
                soft_score + 
                experience_score + 
                education_score + 
                achievement_bonus
            )
            
            # Normalize to 0-100 range
            total_score = min(max(total_score, 0), 100)
            
            # Generate improvement plan
            missing_technical_skills = list(technical_job - technical_resume)
            improvement_plan = self.get_skill_improvement_plan(missing_technical_skills)
            
            analysis = {
                "match_percentage": round(total_score, 1),
                "technical_skills_matched": technical_match,
                "technical_skills_total": technical_total,
                "soft_skills_matched": soft_match,
                "soft_skills_total": soft_total,
                "experience_years": experience_years,
                "education_level": education_level,
                "achievements_count": len(achievements),
                "resume_skills": resume_skills,
                "job_skills": job_skills,
                "keyword_score": round(keyword_score, 1),
                "achievements": achievements,
                "missing_technical_skills": missing_technical_skills,
                "improvement_plan": improvement_plan,
                "generated_projects": generated_projects,
                "matching_skills": list(technical_resume.intersection(technical_job)),
                "detailed_metrics": {
                    "technical_match_rate": f"{(tech_match_rate)*100:.1f}%",
                    "soft_match_rate": f"{(soft_match_rate)*100:.1f}%",
                    "keyword_similarity": f"{keyword_score:.1f}%",
                    "technical_score_contribution": f"{technical_score:.1f}",
                    "keyword_score_contribution": f"{keyword_score_weighted:.1f}",
                    "experience_score_contribution": f"{experience_score:.1f}"
                },
                "score_breakdown": {
                    "Technical Skills": round(technical_score, 1),
                    "Keywords": round(keyword_score_weighted, 1),
                    "Soft Skills": round(soft_score, 1),
                    "Experience": round(experience_score, 1),
                    "Education": round(education_score, 1),
                    "Achievements": round(achievement_bonus, 1)
                }
            }
            
            # Generate profile summary
            analysis['profile_summary'] = self.generate_profile_summary(analysis, resume_text)
            
            return analysis
            
        except Exception as e:
            st.error(f"Analysis error: {str(e)}")
            return self._create_empty_analysis()

    def _create_empty_analysis(self):
        return {
            "match_percentage": 0,
            "technical_skills_matched": 0,
            "technical_skills_total": 0,
            "soft_skills_matched": 0,
            "soft_skills_total": 0,
            "experience_years": 0,
            "education_level": "not_specified",
            "achievements_count": 0,
            "resume_skills": {'technical': [], 'soft': [], 'by_category': {}},
            "job_skills": {'technical': [], 'soft': [], 'by_category': {}},
            "keyword_score": 0,
            "achievements": [],
            "missing_technical_skills": [],
            "improvement_plan": {},
            "generated_projects": [],
            "matching_skills": [],
            "detailed_metrics": {},
            "score_breakdown": {},
            "profile_summary": {
                'strength_level': '',
                'strength_class': '',
                'key_strengths': [],
                'areas_for_improvement': [],
                'competitive_advantage': [],
                'recommendations': [],
                'readiness_score': 0,
                'profile_highlights': {
                    'technical_coverage': 0,
                    'experience_level': '',
                    'achievement_quality': '',
                    'overall_readiness': ''
                }
            }
        }

    def generate_personalized_skill_plan(self, skill, time_available_weeks=1):
        """Generate personalized learning plan based on available time"""
        plan = self.skill_improvement_guides.get(skill, self.default_improvement_plan)
        
        # For one-week intensive plan
        timeline_key = 'learning_path'
        time_mode = '🚀 Intensive Mode (1-Week Fast Track)'
        
        personalized_plan = {
            'skill': skill,
            'time_mode': time_mode,
            'priority': plan.get('priority', 'Medium'),
            'estimated_time': '1 week',
            'learning_path': plan.get(timeline_key, plan.get('learning_path', [])),
            'resources': plan.get('resources', []),
            'practice_platforms': plan.get('practice_platforms', []),
            'project_ideas': plan.get('project_ideas', []),
            'quick_wins': plan.get('quick_wins', []),
            'daily_schedule': self._generate_daily_schedule(skill, time_available_weeks)
        }
        
        return personalized_plan

    def _generate_daily_schedule(self, skill, time_available_weeks):
        """Generate daily study schedule for one-week intensive"""
        return {
            'duration': '6-8 hours/day',
            'breakdown': [
                '📚 Morning (3-4 hours): Theory and concepts',
                '💻 Afternoon (3-4 hours): Hands-on practice and projects',
                '📝 Evening (1-2 hours): Review and planning'
            ],
            'tips': [
                'Take 10-minute breaks every hour',
                'Use Pomodoro technique for focus',
                'Join Discord/Slack communities for quick help',
                'Code along with tutorials for better retention',
                'Build something every day, no matter how small'
            ]
        }
class EnhancedResumeGenerator:
    """Enhanced resume generator with missing skills included in technical skills and project suggestions"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self.technical_skills = {
            'Programming Languages': [
                'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust',
                'swift', 'kotlin', 'scala', 'r', 'matlab', 'perl', 'dart', 'objective-c', 'shell', 'bash',
                'powershell', 'html', 'css', 'sql'
            ],
            'Frameworks & Libraries': [
                'react', 'angular', 'vue', 'django', 'flask', 'spring', 'laravel', 'express', 'node.js',
                'react native', 'flutter', 'tensorflow', 'pytorch', 'pandas', 'numpy', 'scikit-learn',
                'bootstrap', 'jquery', 'asp.net', 'ruby on rails', 'keras', 'opencv'
            ],
            'Databases': [
                'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'oracle', 'sql server', 'sqlite',
                'cassandra', 'dynamodb', 'firebase', 'firestore', 'cosmos db', 'bigquery', 'snowflake'
            ],
            'Cloud & DevOps': [
                'aws', 'amazon web services', 'azure', 'google cloud', 'gcp', 'docker', 'kubernetes',
                'jenkins', 'terraform', 'ansible', 'git', 'github', 'gitlab', 'ci/cd', 'linux', 'unix',
                'nginx', 'apache', 'prometheus', 'grafana'
            ],
            'Data Science & AI': [
                'machine learning', 'deep learning', 'data analysis', 'data visualization', 'nlp',
                'natural language processing', 'computer vision', 'tableau', 'power bi', 'spark', 'hadoop',
                'data mining', 'statistical analysis', 'business intelligence'
            ],
            'Tools & Platforms': [
                'jira', 'confluence', 'slack', 'teams', 'microsoft office', 'excel', 'word', 'powerpoint',
                'visual studio', 'vs code', 'intellij', 'eclipse', 'postman', 'figma', 'sketch', 'photoshop'
            ]
        }
        self.setup_custom_styles()
    
    def setup_custom_styles(self):
        """Setup custom styles for professional resume templates"""
        # Title Style
        self.styles.add(ParagraphStyle(
            name='ResumeTitle',
            parent=self.styles['Title'],
            fontSize=20,
            textColor=colors.HexColor('#5E35B1'),
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        ))
        
        # Section Header Style
        self.styles.add(ParagraphStyle(
            name='ResumeSectionHeader',
            parent=self.styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#7E57C2'),
            spaceBefore=12,
            spaceAfter=6,
            borderBottom=1,
            borderColor=colors.HexColor('#B39DDB'),
            fontName='Helvetica-Bold'
        ))
        
        # Normal Text Style
        self.styles.add(ParagraphStyle(
            name='ResumeNormal',
            parent=self.styles['Normal'],
            fontSize=10,
            leading=13,
            textColor=colors.black
        ))
        
        # Bullet Point Style
        self.styles.add(ParagraphStyle(
            name='ResumeBullet',
            parent=self.styles['Normal'],
            fontSize=10,
            leftIndent=10,
            spaceBefore=3,
            spaceAfter=3
        ))
        
        # Project Bullet Style
        self.styles.add(ParagraphStyle(
            name='ProjectBullet',
            parent=self.styles['Normal'],
            fontSize=10,
            leftIndent=20,
            spaceBefore=3,
            spaceAfter=3,
            bulletIndent=10
        ))
        
        # Highlight Style for Missing Skills
        self.styles.add(ParagraphStyle(
            name='ResumeHighlight',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#FF6B35'),
            leftIndent=10,
            spaceBefore=3,
            spaceAfter=3
        ))

    def extract_personal_info(self, resume_text):
        """Extract personal information from resume text"""
        info = {}
        
        # Extract name (usually at the beginning)
        name_pattern = r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)'
        name_match = re.search(name_pattern, resume_text)
        if name_match:
            info['name'] = name_match.group(1).strip()
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, resume_text)
        if email_match:
            info['email'] = email_match.group(0)
        
        # Extract phone number
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phone_match = re.search(phone_pattern, resume_text)
        if phone_match:
            info['phone'] = phone_match.group(0)
        
        # Extract location
        location_pattern = r'([A-Z][a-z]+(?:[\s-][A-Z][a-z]+)*)(?:\s*,\s*)([A-Z]{2})?'
        location_match = re.search(location_pattern, resume_text[:500])
        if location_match:
            info['location'] = location_match.group(0)
        
        # Extract education
        education_info = self.extract_education_info(resume_text)
        info['education'] = education_info
        
        return info

    def extract_education_info(self, resume_text):
        """Extract detailed education information from resume text"""
        education_info = {}
        
        # Look for education section patterns
        education_sections = re.findall(r'(?i)(education|academic background|qualifications?)[\s\S]*?(?=(experience|skills|projects|$))', resume_text)
        
        if education_sections:
            education_text = education_sections[0][1] if education_sections else ""
            
            # Extract degree information
            degree_patterns = [
                r'(?:bachelor|b\.?s?\.?|b\.?tech|b\.?e\.?)[\s\w]*?(?:in|of)[\s]*([\w\s&,]+)',
                r'(?:master|m\.?s?\.?|m\.?tech|m\.?e\.?)[\s\w]*?(?:in|of)[\s]*([\w\s&,]+)',
                r'(?:ph\.?d|doctorate)[\s\w]*?(?:in|of)[\s]*([\w\s&,]+)'
            ]
            
            for pattern in degree_patterns:
                match = re.search(pattern, education_text, re.IGNORECASE)
                if match:
                    education_info['degree'] = match.group(1).strip()
                    break
            
            # Extract university/college
            university_pattern = r'([A-Z][\w\s&]+(?:University|College|Institute|School))'
            university_match = re.search(university_pattern, education_text)
            if university_match:
                education_info['institution'] = university_match.group(1).strip()
            
            # Extract graduation year
            year_pattern = r'(?:20\d{2}|19\d{2})'
            year_match = re.search(year_pattern, education_text)
            if year_match:
                education_info['year'] = year_match.group(0)
        
        return education_info

    def generate_enhanced_resume(self, analysis, resume_text, job_description, user_info=None, 
                               include_projects=True):
        """Generate enhanced PDF resume with missing skills in technical section and project suggestions"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=72)
        
        story = []
        
        # Use provided personal info or extract from resume
        if not user_info:
            user_info = self.extract_personal_info(resume_text)
        
        # Header Section
        story.extend(self._create_header_section(user_info))
        
        # Professional Summary - Optimized for job description
        story.extend(self._create_enhanced_summary(analysis, job_description))
        
        # Technical Skills - Includes both existing and missing skills
        story.extend(self._create_enhanced_skills_section(analysis))
        
        # Project Suggestions (include at least 2 projects)
        if include_projects:
            story.extend(self._create_project_suggestions_section(analysis))
        
        # Professional Experience - Enhanced with job keywords
        story.extend(self._create_enhanced_experience_section(analysis, resume_text))
        
        # Education Section - Use actual user education
        story.extend(self._create_education_section(analysis, user_info))
        
        # Achievements Section
        if analysis.get('achievements'):
            story.extend(self._create_achievements_section(analysis))
        
        try:
            doc.build(story)
            buffer.seek(0)
            return buffer
        except Exception as e:
            st.error(f"Error generating PDF resume: {str(e)}")
            return None

    def generate_word_resume(self, analysis, resume_text, job_description, user_info=None,
                           include_projects=True):
        """Generate enhanced Word resume with missing skills in technical section and project suggestions"""
        try:
            doc = DocxDocument()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            # Use provided personal info or extract from resume
            if not user_info:
                user_info = self.extract_personal_info(resume_text)
            
            # Header Section
            self._add_word_header(doc, user_info)
            
            # Professional Summary
            self._add_word_summary(doc, analysis, job_description)
            
            # Technical Skills - Includes both existing and missing skills
            self._add_word_skills(doc, analysis)
            
            # Project Suggestions
            if include_projects:
                self._add_word_projects(doc, analysis)
            
            # Professional Experience
            self._add_word_experience(doc, analysis, resume_text)
            
            # Education - Use actual user education
            self._add_word_education(doc, analysis, user_info)
            
            # Achievements
            if analysis.get('achievements'):
                self._add_word_achievements(doc, analysis)
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            st.error(f"Error generating Word resume: {str(e)}")
            return None

    # PDF generation helper methods
    def _create_header_section(self, user_info):
        """Create resume header section"""
        elements = []
        
        # Name
        if user_info.get('name'):
            elements.append(Paragraph(user_info['name'], self.styles['ResumeTitle']))
        else:
            elements.append(Paragraph("Your Name", self.styles['ResumeTitle']))
        
        # Contact Information
        contact_parts = []
        if user_info.get('email'):
            contact_parts.append(user_info['email'])
        if user_info.get('phone'):
            contact_parts.append(user_info['phone'])
        if user_info.get('location'):
            contact_parts.append(user_info['location'])
        
        if contact_parts:
            contact_text = " | ".join(contact_parts)
            elements.append(Paragraph(contact_text, self.styles['ResumeNormal']))
        
        elements.append(Spacer(1, 0.25*inch))
        return elements

    def _create_enhanced_summary(self, analysis, job_description):
        """Create enhanced professional summary"""
        elements = []
        elements.append(Paragraph("PROFESSIONAL SUMMARY", self.styles['ResumeSectionHeader']))
        
        # Build comprehensive summary
        experience = analysis.get('experience_years', 0)
        skills_matched = analysis.get('technical_skills_matched', 0)
        skills_total = analysis.get('technical_skills_total', 1)
        match_percentage = analysis.get('match_percentage', 0)
        
        summary_parts = []
        
        if experience > 0:
            summary_parts.append(f"Results-driven professional with {experience}+ years of experience")
        else:
            summary_parts.append("Motivated professional")
        
        if skills_matched > 0:
            summary_parts.append(f"proficient in {skills_matched} out of {skills_total} required technical skills")
        
        if match_percentage >= 70:
            summary_parts.append("demonstrating strong alignment with position requirements")
        elif match_percentage >= 50:
            summary_parts.append("with solid foundation for role requirements")
        else:
            summary_parts.append("committed to rapid skill development and professional growth")
        
        summary_text = " ".join(summary_parts) + ". "
        summary_text += "Adept at collaborating with cross-functional teams to drive innovation and deliver measurable results through continuous learning and skill enhancement."
        
        elements.append(Paragraph(summary_text, self.styles['ResumeNormal']))
        elements.append(Spacer(1, 0.2*inch))
        return elements

    def _create_enhanced_skills_section(self, analysis):
        """Create enhanced skills section that includes missing skills with development indicators"""
        elements = []
        elements.append(Paragraph("TECHNICAL SKILLS", self.styles['ResumeSectionHeader']))
        
        resume_skills = analysis.get('resume_skills', {})
        matching_skills = analysis.get('matching_skills', [])
        missing_skills = analysis.get('missing_technical_skills', [])
        
        # Combine existing and missing skills (missing skills marked as developing)
        all_skills = {}
        
        # Add existing skills
        skill_categories = resume_skills.get('by_category', {})
        for category, skills in skill_categories.items():
            all_skills[category] = [(skill, 'existing') for skill in skills]
        
        # Add missing skills to appropriate categories or create new category
        for missing_skill in missing_skills:
            category_found = False
            for category in self.technical_skills.keys():
                if missing_skill in self.technical_skills[category]:
                    if category not in all_skills:
                        all_skills[category] = []
                    all_skills[category].append((missing_skill, 'developing'))
                    category_found = True
                    break
            
            # If no category found, add to "Additional Skills"
            if not category_found:
                if "Additional Skills" not in all_skills:
                    all_skills["Additional Skills"] = []
                all_skills["Additional Skills"].append((missing_skill, 'developing'))
        
        # Display skills with indicators
        if all_skills:
            for category, skills in all_skills.items():
                if skills:
                    # Format skills with development indicators
                    skill_list = []
                    for skill, status in skills:
                        if status == 'existing':
                            skill_list.append(skill.title())
                        else:
                            skill_list.append(f"{skill.title()} (Developing)")
                    
                    skill_text = f"<b>{category}:</b> {', '.join(skill_list)}"
                    elements.append(Paragraph(skill_text, self.styles['ResumeNormal']))
        
        # Add soft skills
        if resume_skills.get('soft'):
            soft_skills_text = f"<b>Soft Skills:</b> {', '.join(resume_skills['soft'])}"
            elements.append(Paragraph(soft_skills_text, self.styles['ResumeNormal']))
        
        elements.append(Spacer(1, 0.2*inch))
        return elements

    def _create_project_suggestions_section(self, analysis):
        """Create project suggestions section with at least 2 projects"""
        elements = []
        elements.append(Paragraph("PROJECTS", self.styles['ResumeSectionHeader']))
        
        generated_projects = analysis.get('generated_projects', [])
        
        # Ensure we have at least 2 projects
        if len(generated_projects) < 2:
            # Add default projects if needed
            default_projects = self._get_default_projects(analysis)
            generated_projects.extend(default_projects[:2 - len(generated_projects)])
        
        if generated_projects:
            for i, project in enumerate(generated_projects[:3]):  # Show up to 3 projects
                # Project title with skills
                project_title = f"<b>{project['name']}</b>"
                elements.append(Paragraph(project_title, self.styles['ResumeNormal']))
                
                # Skills used
                skills_text = f"<i>Technologies: {', '.join(project['skills'])}</i>"
                elements.append(Paragraph(skills_text, self.styles['ResumeNormal']))
                
                # Project bullets with impact metrics
                for bullet in project['bullets'][:3]:  # Show up to 3 bullets per project
                    elements.append(Paragraph(f"• {bullet}", self.styles['ProjectBullet']))
                
                elements.append(Spacer(1, 0.1*inch))
        else:
            elements.append(Paragraph("No specific projects generated.", self.styles['ResumeNormal']))
        
        elements.append(Spacer(1, 0.2*inch))
        return elements

    def _get_default_projects(self, analysis):
        """Generate default projects based on skills analysis"""
        default_projects = []
        
        # Get skills for project generation
        matching_skills = analysis.get('matching_skills', [])
        missing_skills = analysis.get('missing_technical_skills', [])
        
        # Project 1: Web Development focused
        if any(skill in ['javascript', 'python', 'react', 'node.js'] for skill in matching_skills + missing_skills):
            default_projects.append({
                'name': 'Full-Stack Web Application',
                'skills': ['React', 'Node.js', 'MongoDB', 'Express'],
                'bullets': [
                    'Developed responsive web application with user authentication and real-time features',
                    'Implemented RESTful API with proper error handling and security measures',
                    'Deployed application using cloud services with CI/CD pipeline integration'
                ]
            })
        
        # Project 2: Data Analysis focused
        if any(skill in ['python', 'pandas', 'sql', 'data analysis'] for skill in matching_skills + missing_skills):
            default_projects.append({
                'name': 'Data Analysis Dashboard',
                'skills': ['Python', 'Pandas', 'SQL', 'Data Visualization'],
                'bullets': [
                    'Analyzed large datasets to extract meaningful insights and trends',
                    'Created interactive visualizations and dashboards for data presentation',
                    'Automated data processing pipelines reducing manual effort by 70%'
                ]
            })
        
        # Project 3: General software development
        if len(default_projects) < 2:
            default_projects.append({
                'name': 'Portfolio Website & Blog',
                'skills': ['HTML', 'CSS', 'JavaScript', 'Responsive Design'],
                'bullets': [
                    'Designed and developed personal portfolio website with blog functionality',
                    'Optimized for performance and SEO, achieving 95+ Google PageSpeed score',
                    'Implemented responsive design ensuring compatibility across all devices'
                ]
            })
        
        return default_projects[:2]  # Return exactly 2 projects

    def _create_education_section(self, analysis, user_info):
        """Create education section with user's actual education information"""
        elements = []
        elements.append(Paragraph("EDUCATION", self.styles['ResumeSectionHeader']))
        
        # Use extracted education info if available
        if user_info.get('education'):
            edu_info = user_info['education']
            degree = edu_info.get('degree', 'Degree in Computer Science or related field')
            institution = edu_info.get('institution', 'University/College')
            year = edu_info.get('year', 'Graduation Year')
            
            education_text = f"<b>{degree}</b><br/>{institution}"
            if year:
                education_text += f" | {year}"
        else:
            # Fallback to analysis-based education
            education_level = analysis.get('education_level', 'not_specified')
            if education_level != 'not_specified':
                education_text = f"<b>{education_level.title()} Degree</b> - Computer Science or related field"
            else:
                education_text = "<b>Degree</b> - Computer Science or related field"
        
        elements.append(Paragraph(education_text, self.styles['ResumeNormal']))
        elements.append(Spacer(1, 0.2*inch))
        return elements

    def _create_enhanced_experience_section(self, analysis, resume_text):
        """Create enhanced experience section"""
        elements = []
        elements.append(Paragraph("PROFESSIONAL EXPERIENCE", self.styles['ResumeSectionHeader']))
        
        experience_years = analysis.get('experience_years', 0)
        if experience_years > 0:
            exp_text = f"Seasoned professional with {experience_years}+ years of comprehensive experience in software development and technical implementation. Demonstrated success in delivering scalable solutions and driving business growth through technical innovation."
        else:
            exp_text = "Enthusiastic professional seeking opportunity to apply technical skills and learning capabilities. Quick learner with strong problem-solving abilities and commitment to continuous improvement."
        
        elements.append(Paragraph(exp_text, self.styles['ResumeNormal']))
        
        # Add achievements if available
        if analysis.get('achievements'):
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Key Achievements:", self.styles['ResumeNormal']))
            for achievement in analysis['achievements'][:3]:
                elements.append(Paragraph(f"• {achievement}", self.styles['ResumeBullet']))
        
        elements.append(Spacer(1, 0.2*inch))
        return elements

    def _create_achievements_section(self, analysis):
        """Create achievements section"""
        elements = []
        elements.append(Paragraph("KEY ACHIEVEMENTS", self.styles['ResumeSectionHeader']))
        
        for achievement in analysis.get('achievements', [])[:5]:
            elements.append(Paragraph(f"• {achievement}", self.styles['ResumeBullet']))
        
        elements.append(Spacer(1, 0.2*inch))
        return elements

    # Word document methods
    def _add_word_header(self, doc, user_info):
        """Add header to Word document"""
        if user_info.get('name'):
            doc.add_heading(user_info['name'], 0)
        else:
            doc.add_heading("Your Name", 0)
        
        contact_info = []
        if user_info.get('email'):
            contact_info.append(user_info['email'])
        if user_info.get('phone'):
            contact_info.append(user_info['phone'])
        if user_info.get('location'):
            contact_info.append(user_info['location'])
        
        if contact_info:
            doc.add_paragraph(' | '.join(contact_info))
        
        doc.add_paragraph()

    def _add_word_summary(self, doc, analysis, job_description):
        """Add summary to Word document"""
        doc.add_heading('Professional Summary', level=1)
        
        experience = analysis.get('experience_years', 0)
        skills_matched = analysis.get('technical_skills_matched', 0)
        skills_total = analysis.get('technical_skills_total', 1)
        
        summary = f"Results-driven professional with {experience}+ years of experience "
        summary += f"proficient in {skills_matched} out of {skills_total} required technical skills. "
        summary += "Demonstrated ability to deliver measurable results through continuous learning and skill enhancement."
        
        doc.add_paragraph(summary)
        doc.add_paragraph()

    def _add_word_skills(self, doc, analysis):
        """Add skills to Word document including missing skills with development indicators"""
        doc.add_heading('Technical Skills', level=1)
        
        resume_skills = analysis.get('resume_skills', {})
        matching_skills = analysis.get('matching_skills', [])
        missing_skills = analysis.get('missing_technical_skills', [])
        
        # Combine existing and missing skills
        all_skills = {}
        
        # Add existing skills
        for category, skills in resume_skills.get('by_category', {}).items():
            all_skills[category] = [(skill, 'existing') for skill in skills]
        
        # Add missing skills to appropriate categories
        for missing_skill in missing_skills:
            category_found = False
            for category in self.technical_skills.keys():
                if missing_skill in self.technical_skills[category]:
                    if category not in all_skills:
                        all_skills[category] = []
                    all_skills[category].append((missing_skill, 'developing'))
                    category_found = True
                    break
            
            if not category_found:
                if "Additional Skills" not in all_skills:
                    all_skills["Additional Skills"] = []
                all_skills["Additional Skills"].append((missing_skill, 'developing'))
        
        # Display skills with development indicators
        for category, skills in all_skills.items():
            skill_list = []
            for skill, status in skills:
                if status == 'existing':
                    skill_list.append(skill.title())
                else:
                    skill_list.append(f"{skill.title()} (Developing)")
            
            doc.add_paragraph(f"{category}: {', '.join(skill_list)}")
        
        if resume_skills.get('soft'):
            doc.add_paragraph(f"Soft Skills: {', '.join(resume_skills['soft'])}")
        
        doc.add_paragraph()

    def _add_word_projects(self, doc, analysis):
        """Add project suggestions to Word document with at least 2 projects"""
        doc.add_heading('Projects', level=1)
        
        generated_projects = analysis.get('generated_projects', [])
        
        # Ensure we have at least 2 projects
        if len(generated_projects) < 2:
            default_projects = self._get_default_projects(analysis)
            generated_projects.extend(default_projects[:2 - len(generated_projects)])
        
        for project in generated_projects[:3]:  # Show up to 3 projects
            # Project title
            p = doc.add_paragraph()
            p.add_run(f"{project['name']}").bold = True
            
            # Skills used
            doc.add_paragraph(f"Technologies: {', '.join(project['skills'])}")
            
            # Project bullets
            for bullet in project['bullets'][:3]:
                doc.add_paragraph(f"• {bullet}", style='List Bullet')
            
            doc.add_paragraph()
        
        doc.add_paragraph()

    def _add_word_experience(self, doc, analysis, resume_text):
        """Add experience to Word document"""
        doc.add_heading('Professional Experience', level=1)
        
        experience_years = analysis.get('experience_years', 0)
        if experience_years > 0:
            doc.add_paragraph(f"Professional with {experience_years}+ years of comprehensive experience")
            doc.add_paragraph("• Demonstrated expertise in key technical domains with proven track record", style='List Bullet')
            doc.add_paragraph("• Successfully collaborated with cross-functional teams", style='List Bullet')
            
            # Add achievements if available
            if analysis.get('achievements'):
                doc.add_paragraph("Key Achievements:")
                for achievement in analysis['achievements'][:3]:
                    doc.add_paragraph(f"• {achievement}", style='List Bullet 2')
        else:
            doc.add_paragraph("Seeking opportunity to apply technical skills and learning capabilities")
            doc.add_paragraph("• Quick learner with strong problem-solving abilities", style='List Bullet')
        
        doc.add_paragraph()

    def _add_word_education(self, doc, analysis, user_info):
        """Add education to Word document with user's actual education"""
        doc.add_heading('Education', level=1)
        
        # Use extracted education info if available
        if user_info.get('education'):
            edu_info = user_info['education']
            degree = edu_info.get('degree', 'Degree in Computer Science or related field')
            institution = edu_info.get('institution', 'University/College')
            year = edu_info.get('year', 'Graduation Year')
            
            education_text = f"{degree} | {institution}"
            if year:
                education_text += f" | {year}"
        else:
            # Fallback to analysis-based education
            education_level = analysis.get('education_level', 'not_specified')
            if education_level != 'not_specified':
                education_text = f"{education_level.title()} Degree - Computer Science or related field"
            else:
                education_text = "Degree - Computer Science or related field"
        
        doc.add_paragraph(education_text)
        doc.add_paragraph()

    def _add_word_achievements(self, doc, analysis):
        """Add achievements to Word document"""
        doc.add_heading('Key Achievements', level=1)
        
        for achievement in analysis.get('achievements', [])[:5]:
            doc.add_paragraph(f"• {achievement}", style='List Bullet')
        
        doc.add_paragraph()
    # Database Functions
def init_db():
    conn = sqlite3.connect("DB_PATH", check_same_thread=False)
    c = conn.cursor()
    
    c.execute("""
        CREATE TABLE IF NOT EXISTS users (
            username TEXT PRIMARY KEY,
            password_hash BLOB NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    
    c.execute("""
        CREATE TABLE IF NOT EXISTS analysis_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT,
            job_title TEXT,
            match_percentage REAL,
            analysis_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            analysis_result TEXT
        )
    """)
    
    conn.commit()
    conn.close()

def add_user(username, password):
    conn = sqlite3.connect("DB_PATH", check_same_thread=False)
    c = conn.cursor()
    password_hash = bcrypt.hashpw(password.encode(), bcrypt.gensalt())
    try:
        c.execute("INSERT INTO users (username, password_hash) VALUES (?, ?)", 
                 (username, password_hash))
        conn.commit()
        return True
    except sqlite3.IntegrityError:
        return False
    finally:
        conn.close()

def verify_user(username, password):
    conn = sqlite3.connect("DB_PATH", check_same_thread=False)
    c = conn.cursor()
    c.execute("SELECT password_hash FROM users WHERE username = ?", (username,))
    row = c.fetchone()
    if row:
        stored_hash = row[0]
        if bcrypt.checkpw(password.encode(), stored_hash):
            conn.close()
            return True
    conn.close()
    return False

def save_analysis(username, job_title, match_percentage, analysis_result):
    conn = sqlite3.connect("DB_PATH", check_same_thread=False)
    c = conn.cursor()
    c.execute("""
        INSERT INTO analysis_history (username, job_title, match_percentage, analysis_result)
        VALUES (?, ?, ?, ?)
    """, (username, job_title, match_percentage, json.dumps(analysis_result)))
    conn.commit()
    conn.close()

def get_user_history(username):
    conn = sqlite3.connect("DB_PATH", check_same_thread=False)
    c = conn.cursor()
    c.execute("""
        SELECT id, job_title, match_percentage, analysis_date, analysis_result
        FROM analysis_history 
        WHERE username = ? 
        ORDER BY analysis_date DESC 
        LIMIT 10
    """, (username,))
    history = c.fetchall()
    conn.close()
    return history

def get_analysis_by_id(analysis_id):
    conn = sqlite3.connect("DB_PATH", check_same_thread=False)
    c = conn.cursor()
    c.execute("SELECT analysis_result FROM analysis_history WHERE id = ?", (analysis_id,))
    row = c.fetchone()
    conn.close()
    if row:
        return json.loads(row[0])
    return None

# File processing
def extract_text_from_file(uploaded_file):
    """Extract text from PDF or DOCX files"""
    try:
        if uploaded_file.type == "application/pdf":
            reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text
        elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            doc = Document(uploaded_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        else:
            st.error("Unsupported file format. Please upload PDF or DOCX.")
            return ""
    except Exception as e:
        st.error(f"Error reading file: {str(e)}")
        return ""

# Enhanced Visualization functions
def create_advanced_match_gauge(percentage):
    """Create enhanced gauge chart with more details"""
    if percentage >= 85:
        delta_ref = 85
        delta_color = "#66BB6A"
        bar_color = "#66BB6A"
    elif percentage >= 70:
        delta_ref = 70
        delta_color = "#26A69A"
        bar_color = "#26A69A"
    elif percentage >= 50:
        delta_ref = 50
        delta_color = "#FFA726"
        bar_color = "#FFA726"
    else:
        delta_ref = 30
        delta_color = "#EF5350"
        bar_color = "#EF5350"
    
    fig = go.Figure(go.Indicator(
        mode = "gauge+number+delta",
        value = percentage,
        domain = {'x': [0, 1], 'y': [0, 1]},
        title = {'text': "ATS Match Score", 'font': {'size': 26, 'family': 'Inter', 'color': '#2E2E3A'}},
        delta = {
            'reference': delta_ref, 
            'increasing': {'color': delta_color},
            'font': {'size': 20}
        },
        number = {'font': {'size': 50, 'color': bar_color}},
        gauge = {
            'axis': {'range': [None, 100], 'tickwidth': 2, 'tickcolor': "#666676", 'tickfont': {'size': 14}},
            'bar': {'color': bar_color, 'thickness': 0.8},
            'bgcolor': "white",
            'borderwidth': 3,
            'bordercolor': "#E1E1E8",
            'steps': [
                {'range': [0, 50], 'color': "rgba(239, 83, 80, 0.1)"},
                {'range': [50, 70], 'color': "rgba(255, 167, 38, 0.1)"},
                {'range': [70, 85], 'color': "rgba(38, 166, 154, 0.1)"},
                {'range': [85, 100], 'color': "rgba(102, 187, 106, 0.1)"}
            ],
            'threshold': {
                'line': {'color': "red", 'width': 4},
                'thickness': 0.75,
                'value': 80
            }
        }
    ))
    
    fig.update_layout(
        height=350,
        font={'family': "Inter"},
        paper_bgcolor="rgba(255,255,255,0.9)",
        margin=dict(l=20, r=20, t=60, b=20)
    )
    return fig

def create_skill_match_chart(technical_match, technical_total, soft_match, soft_total):
    """Enhanced skill match visualization"""
    categories = ['Technical Skills', 'Soft Skills']
    matched = [technical_match, soft_match]
    required = [technical_total, soft_total]
    match_rates = [(matched[i]/max(required[i], 1))*100 for i in range(2)]
    
    fig = go.Figure()
    
    # Add matched bars
    fig.add_trace(go.Bar(
        name='Your Skills', 
        x=categories, 
        y=matched,
        text=[f'{matched[i]}<br>({match_rates[i]:.0f}%)' for i in range(2)],
        textposition='auto',
        marker_color='#7E57C2',
        hovertemplate='<b>%{x}</b><br>Matched: %{y}<br>Rate: %{text}<extra></extra>'
    ))
    
    # Add required line
    fig.add_trace(go.Scatter(
        name='Required',
        x=categories,
        y=required,
        mode='lines+markers+text',
        text=[f'Target: {r}' for r in required],
        textposition='top center',
        line=dict(color='#EF5350', width=4, dash='dash'),
        marker=dict(size=12, color='#EF5350', symbol='diamond')
    ))
    
    fig.update_layout(
        title={
            'text': "Skill Match Analysis",
            'font': {'size': 20, 'family': 'Inter', 'color': '#2E2E3A'}
        },
        barmode='overlay',
        height=450,
        showlegend=True,
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="right",
            x=1
        ),
        plot_bgcolor='rgba(255,255,255,0.9)',
        paper_bgcolor='rgba(255,255,255,0.9)',
        font={'family': 'Inter'},
        xaxis=dict(
            showgrid=False,
            title_font={'size': 14}
        ),
        yaxis=dict(
            showgrid=True,
            gridcolor='rgba(126, 87, 194, 0.1)',
            title='Number of Skills',
            title_font={'size': 14}
        )
    )
    return fig

def create_score_breakdown_chart(score_breakdown):
    """Create detailed score breakdown visualization"""
    categories = list(score_breakdown.keys())
    values = list(score_breakdown.values())
    
    colors = ['#7E57C2', '#26A69A', '#66BB6A', '#FFA726', '#42A5F5', '#AB47BC']
    
    fig = go.Figure(data=[
        go.Bar(
            x=categories,
            y=values,
            text=[f'{v:.1f}' for v in values],
            textposition='auto',
            marker=dict(
                color=colors[:len(categories)],
                line=dict(color='white', width=2)
            ),
            hovertemplate='<b>%{x}</b><br>Score: %{y:.1f}<extra></extra>'
        )
    ])
    
    fig.update_layout(
        title={
            'text': "Detailed Score Breakdown",
            'font': {'size': 20, 'family': 'Inter', 'color': '#2E2E3A'}
        },
        height=400,
        plot_bgcolor='rgba(255,255,255,0.9)',
        paper_bgcolor='rgba(255,255,255,0.9)',
        font={'family': 'Inter'},
        xaxis=dict(
            showgrid=False,
            title_font={'size': 14}
        ),
        yaxis=dict(
            showgrid=True,
            gridcolor='rgba(126, 87, 194, 0.1)',
            title='Score Contribution',
            title_font={'size': 14},
            range=[0, max(values) * 1.2]
        ),
        showlegend=False
    )
    
    return fig

def create_missing_skills_chart(missing_skills, improvement_plan):
    """Enhanced missing skills visualization"""
    if not missing_skills:
        return None
    
    skills = list(missing_skills)[:10]
    priorities = [improvement_plan.get(skill, {}).get('priority', 'Medium') for skill in skills]
    time_estimates = []
    
    for skill in skills:
        plan = improvement_plan.get(skill, {})
        time_est = plan.get('time_estimate', '1 week')
        # Extract numeric value for sorting
        if 'week' in time_est:
            time_estimates.append(1)
        elif 'month' in time_est:
            time_estimates.append(4)
        else:
            time_estimates.append(2)
    
    # Color mapping
    color_map = {'High': '#EF5350', 'Medium': '#FFA726', 'Low': '#7E57C2'}
    colors = [color_map.get(p, '#7E57C2') for p in priorities]
    
    fig = go.Figure(data=[
        go.Bar(
            x=skills, 
            y=time_estimates,
            marker_color=colors,
            text=[f'{p}<br>{t}w' for p, t in zip(priorities, time_estimates)],
            textposition='auto',
            hovertemplate='<b>%{x}</b><br>Priority: %{text}<br>Time: %{y} week(s)<extra></extra>'
        )
    ])
    
    fig.update_layout(
        title={
            'text': "Missing Skills - Learning Time & Priority",
            'font': {'size': 20, 'family': 'Inter', 'color': '#2E2E3A'}
        },
        xaxis_title="Skills",
        yaxis_title="Weeks Required",
        height=450,
        plot_bgcolor='rgba(255,255,255,0.9)',
        paper_bgcolor='rgba(255,255,255,0.9)',
        font={'family': 'Inter'},
        xaxis=dict(
            showgrid=False,
            tickangle=-45,
            title_font={'size': 14}
        ),
        yaxis=dict(
            showgrid=True,
            gridcolor='rgba(126, 87, 194, 0.1)',
            title_font={'size': 14},
            range=[0, 5]
        )
    )
    return fig

def create_radar_chart(analysis):
    """Create radar chart for profile analysis"""
    categories = ['Technical<br>Skills', 'Soft<br>Skills', 'Experience', 'Keywords', 'Achievements']
    
    tech_score = (analysis['technical_skills_matched'] / max(analysis['technical_skills_total'], 1)) * 100
    soft_score = (analysis['soft_skills_matched'] / max(analysis['soft_skills_total'], 1)) * 100
    exp_score = min(analysis['experience_years'] * 20, 100)
    keyword_score = analysis['keyword_score']
    achievement_score = min(analysis['achievements_count'] * 20, 100)
    
    values = [tech_score, soft_score, exp_score, keyword_score, achievement_score]
    
    fig = go.Figure()
    
    fig.add_trace(go.Scatterpolar(
        r=values,
        theta=categories,
        fill='toself',
        fillcolor='rgba(126, 87, 194, 0.3)',
        line=dict(color='#7E57C2', width=3),
        marker=dict(size=8, color='#7E57C2'),
        name='Your Profile',
        hovertemplate='<b>%{theta}</b><br>Score: %{r:.1f}%<extra></extra>'
    ))
    
    # Add benchmark line
    benchmark = [80, 80, 80, 80, 80]
    fig.add_trace(go.Scatterpolar(
        r=benchmark,
        theta=categories,
        fill='toself',
        fillcolor='rgba(102, 187, 106, 0.1)',
        line=dict(color='#66BB6A', width=2, dash='dash'),
        marker=dict(size=6, color='#66BB6A'),
        name='Target (80%)',
        hovertemplate='<b>%{theta}</b><br>Target: %{r:.1f}%<extra></extra>'
    ))
    
    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 100],
                showticklabels=True,
                tickfont=dict(size=12),
                gridcolor='rgba(126, 87, 194, 0.2)'
            ),
            angularaxis=dict(
                gridcolor='rgba(126, 87, 194, 0.2)'
            ),
            bgcolor='rgba(255,255,255,0.9)'
        ),
        showlegend=True,
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=-0.2,
            xanchor="center",
            x=0.5
        ),
        height=500,
        title={
            'text': "Profile Strength Analysis",
            'font': {'size': 20, 'family': 'Inter', 'color': '#2E2E3A'},
            'x': 0.5,
            'xanchor': 'center'
        },
        paper_bgcolor='rgba(255,255,255,0.9)',
        font={'family': 'Inter'}
    )
    
    return fig

def create_progress_timeline(history):
    """Create enhanced progress timeline chart"""
    if len(history) <= 1:
        return None
    
    df = pd.DataFrame(history, columns=['ID', 'Job Title', 'Match %', 'Date', 'Analysis Result'])
    df['Date'] = pd.to_datetime(df['Date']).dt.strftime('%Y-%m-%d')
    
    # Create a more detailed timeline with trend line
    fig = go.Figure()
    
    # Main line with markers
    fig.add_trace(go.Scatter(
        x=df['Date'],
        y=df['Match %'],
        mode='lines+markers+text',
        name='ATS Score',
        line=dict(color='#7E57C2', width=4),
        marker=dict(size=12, color='#7E57C2'),
        text=[f"{score}%" for score in df['Match %']],
        textposition='top center',
        hovertemplate='<b>%{x}</b><br>Score: %{y:.1f}%<extra></extra>'
    ))
    
    # Add trend line
    if len(df) > 2:
        z = np.polyfit(range(len(df)), df['Match %'], 1)
        trend_line = np.poly1d(z)(range(len(df)))
        fig.add_trace(go.Scatter(
            x=df['Date'],
            y=trend_line,
            mode='lines',
            name='Trend',
            line=dict(color='#26A69A', width=2, dash='dash'),
            hovertemplate='Trend: %{y:.1f}%<extra></extra>'
        ))
    
    fig.update_layout(
        title={
            'text': "📈 Progress Timeline & Trend Analysis",
            'font': {'size': 22, 'family': 'Inter', 'color': '#2E2E3A'}
        },
        height=500,
        plot_bgcolor='rgba(255,255,255,0.9)',
        paper_bgcolor='rgba(255,255,255,0.9)',
        font={'family': 'Inter'},
        xaxis=dict(
            showgrid=True,
            gridcolor='rgba(126, 87, 194, 0.1)',
            title='Analysis Date',
            title_font={'size': 14}
        ),
        yaxis=dict(
            showgrid=True,
            gridcolor='rgba(126, 87, 194, 0.1)',
            title='ATS Match Score (%)',
            title_font={'size': 14},
            range=[0, 100]
        ),
        showlegend=True,
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="right",
            x=1
        )
    )
    
    return fig
# Initialize database
init_db()

# Initialize session state
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
if 'username' not in st.session_state:
    st.session_state.username = ""
if 'analyzer' not in st.session_state:
    st.session_state.analyzer = AdvancedResumeAnalyzer()
if 'resume_generator' not in st.session_state:
    st.session_state.resume_generator = EnhancedResumeGenerator()
if 'current_analysis' not in st.session_state:
    st.session_state.current_analysis = None
if 'user_info' not in st.session_state:
    st.session_state.user_info = {}
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = ""
if 'job_description' not in st.session_state:
    st.session_state.job_description = ""

# Main Application
if not st.session_state.logged_in:
    # Hero Section with Brand Name - CENTERED
    st.markdown("""
    <div class="hero-section">
        <div class="hero-content">
            <div style="font-size: 1.2rem; font-weight: 600; opacity: 0.9; margin-bottom: 1rem; letter-spacing: 3px; text-align: center;">WELCOME TO</div>
            <h1 class="hero-title" style="font-size: 4.5rem; margin-bottom: 1.5rem; text-align: center;">
                🚀 <span style="background: linear-gradient(135deg, #FFFFFF 0%, #E1E1E8 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent;">AscendCV</span>
            </h1>
            <p class="hero-subtitle" style="font-size: 1.4rem; margin-bottom: 2.5rem; line-height: 1.8; text-align: center;">
                Your AI-Powered Resume Analyzer Platform<br/>
                <span style="opacity: 0.85; font-size: 1.1rem; text-align: center;">Optimize Your Resume • Match Job Requirements • Accelerate Skill Development</span>
            </p>
            <div class="hero-features" style="text-align: center;">
                <div class="feature-badge">🚀 Instant ATS Analysis</div>
                <div class="feature-badge">🎯 Personalized Learning</div>
                <div class="feature-badge">📊 Career Analytics</div>
                <div class="feature-badge">💼 Smart Recommendations</div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Main Content Container - Two Column Layout
    st.markdown("""
    <div style="max-width: 1400px; margin: 0 auto; padding: 0 2rem;">
    """, unsafe_allow_html=True)
    
    # Create main two columns: Login/Signup (left) and Features (right)
    main_col1, main_col2 = st.columns([1, 1.3], gap="large")
    
    # LEFT SIDE - Login/Signup Section
    with main_col1:
        st.markdown("""
        <div style="position: sticky; top: 2rem;">
        """, unsafe_allow_html=True)
        
        # Initialize session state for form visibility
        if 'show_login' not in st.session_state:
            st.session_state.show_login = False
        if 'show_signup' not in st.session_state:
            st.session_state.show_signup = False
        
        # Main container with welcome message
        st.markdown("""
        <div style="background: rgba(255,255,255,0.95); backdrop-filter: blur(20px); border-radius: 2rem; padding: 2.5rem; border: 1px solid rgba(255,255,255,0.5); box-shadow: var(--shadow-lg); margin-bottom: 2rem;">
            <div style="text-align: center; margin-bottom: 1.5rem;">
                <h2 style="color: var(--primary); font-weight: 800; font-size: 2.2rem; margin-bottom: 1rem;">Get Started</h2>
                <p style="color: var(--text-secondary); font-size: 1rem; line-height: 1.6;">Access your personalized AI-powered career optimization dashboard</p>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # Action Buttons
        col_btn1, col_btn2 = st.columns(2, gap="medium")
        
        with col_btn1:
            if st.button("Login", use_container_width=True, type="primary" if not st.session_state.show_login else "secondary"):
                st.session_state.show_login = not st.session_state.show_login
                st.session_state.show_signup = False
                st.rerun()
        
        with col_btn2:
            if st.button("Sign Up", use_container_width=True, type="primary" if not st.session_state.show_signup else "secondary"):
                st.session_state.show_signup = not st.session_state.show_signup
                st.session_state.show_login = False
                st.rerun()
        
        st.markdown('<div style="margin: 1.5rem 0;"></div>', unsafe_allow_html=True)
        
        # Login Form (Show/Hide based on state)
        if st.session_state.show_login:
            st.markdown("""
            <div style="background: linear-gradient(135deg, rgba(126, 87, 194, 0.05), rgba(94, 53, 177, 0.02)); backdrop-filter: blur(20px); border-radius: 1.5rem; padding: 2rem; border: 2px solid rgba(126, 87, 194, 0.3); box-shadow: var(--shadow-md); margin-bottom: 1.5rem; animation: slideIn 0.3s ease-out;">
                <div style="text-align: center; margin-bottom: 1.5rem;">
                    <h3 style="color: var(--primary); font-weight: 700; font-size: 1.5rem; margin-bottom: 0.3rem;">👋 Welcome Back!</h3>
                    <p style="color: var(--text-secondary); font-size: 0.9rem;">Sign in to continue your journey</p>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            with st.form("login_form", clear_on_submit=False):
                username = st.text_input("Username", placeholder="Enter your username", key="login_username", label_visibility="collapsed")
                st.markdown('<div style="margin-bottom: 0.5rem;"></div>', unsafe_allow_html=True)
                password = st.text_input("Password", type="password", placeholder="Enter your password", key="login_password", label_visibility="collapsed")
                st.markdown('<div style="margin-bottom: 1rem;"></div>', unsafe_allow_html=True)
                
                col1, col2 = st.columns([1, 1])
                with col1:
                    login_btn = st.form_submit_button("Login", type="primary", use_container_width=True)
                with col2:
                    cancel_login = st.form_submit_button("Cancel", use_container_width=True)
                
                if cancel_login:
                    st.session_state.show_login = False
                    st.rerun()
                
                if login_btn:
                    if not username or not password:
                        st.error("⚠️ Please fill in all fields")
                    elif verify_user(username, password):
                        st.session_state.logged_in = True
                        st.session_state.username = username
                        st.success("✅ Login successful! Redirecting...")
                        time.sleep(1)
                        st.rerun()
                    else:
                        st.error("❌ Invalid username or password")
        
        # Signup Form (Show/Hide based on state)
        if st.session_state.show_signup:
            st.markdown("""
            <div style="background: linear-gradient(135deg, rgba(38, 166, 154, 0.05), rgba(0, 121, 107, 0.02)); backdrop-filter: blur(20px); border-radius: 1.5rem; padding: 2rem; border: 2px solid rgba(38, 166, 154, 0.3); box-shadow: var(--shadow-md); margin-bottom: 1.5rem; animation: slideIn 0.3s ease-out;">
                <div style="text-align: center; margin-bottom: 1.5rem;">
                    <h3 style="color: var(--secondary); font-weight: 700; font-size: 1.5rem; margin-bottom: 0.3rem;">🎉 Join AscendCV</h3>
                    <p style="color: var(--text-secondary); font-size: 0.9rem;">Create your account in seconds</p>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            with st.form("signup_form", clear_on_submit=False):
                new_username = st.text_input("Username", placeholder="Choose a username", key="signup_username", label_visibility="collapsed")
                st.markdown('<div style="margin-bottom: 0.5rem;"></div>', unsafe_allow_html=True)
                new_password = st.text_input("Password", type="password", placeholder="Create a password (min 6 characters)", key="signup_password", label_visibility="collapsed")
                st.markdown('<div style="margin-bottom: 0.5rem;"></div>', unsafe_allow_html=True)
                confirm_password = st.text_input("Confirm Password", type="password", placeholder="Confirm your password", key="confirm_password", label_visibility="collapsed")
                st.markdown('<div style="margin-bottom: 1rem;"></div>', unsafe_allow_html=True)
                
                col1, col2 = st.columns([1, 1])
                with col1:
                    signup_btn = st.form_submit_button("Create Account", type="primary", use_container_width=True)
                with col2:
                    cancel_signup = st.form_submit_button("Cancel", use_container_width=True)
                
                if cancel_signup:
                    st.session_state.show_signup = False
                    st.rerun()
                
                if signup_btn:
                    if not new_username or not new_password:
                        st.error("⚠️ Please fill in all fields")
                    elif new_password != confirm_password:
                        st.error("❌ Passwords don't match")
                    elif len(new_password) < 6:
                        st.error("⚠️ Password must be at least 6 characters")
                    else:
                        if add_user(new_username, new_password):
                            st.success("✅ Account created successfully! Please login.")
                            st.session_state.show_signup = False
                            st.session_state.show_login = True
                            time.sleep(1)
                            st.rerun()
                        else:
                            st.error("❌ Username already exists. Please choose another.")
        
        # Info box when no form is open
        if not st.session_state.show_login and not st.session_state.show_signup:
            st.markdown("""
            <div style="background: linear-gradient(135deg, rgba(102, 187, 106, 0.08), rgba(76, 175, 80, 0.03)); backdrop-filter: blur(20px); border-radius: 1.5rem; padding: 2rem; border: 1px solid rgba(102, 187, 106, 0.2); box-shadow: var(--shadow-sm); margin-top: 1.5rem;">
                <div style="text-align: center;">
                    <div style="font-size: 2.5rem; margin-bottom: 1rem;">🎯</div>
                    <h4 style="color: var(--text-primary); font-weight: 600; margin-bottom: 0.8rem;">Ready to Transform Your Career?</h4>
                    <p style="color: var(--text-secondary); font-size: 0.9rem; line-height: 1.6;">
                        Click <strong>Login</strong> if you already have an account, or <strong>Sign Up</strong> to create a new one and start optimizing your resume today!
                    </p>
                </div>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    # RIGHT SIDE - Features Section
    with main_col2:
        st.markdown("""
        <div style="margin-bottom: 2rem;">
            <h2 style="font-size: 2.5rem; font-weight: 800; background: var(--gradient-primary); -webkit-background-clip: text; -webkit-text-fill-color: transparent; margin-bottom: 1.5rem; text-align: center;">
                What AscendCV Offers
            </h2>
            <p style="font-size: 1.1rem; color: var(--text-secondary); margin: 0 auto 2.5rem auto; line-height: 1.8; text-align: justify;">
                <strong>AscendCV</strong> uses advanced AI algorithms to analyze your resume against job descriptions, providing instant ATS compatibility scores, 
                identifying skill gaps, and creating personalized learning roadmaps. Our intelligent system compares your qualifications with job requirements, 
                highlights your competitive strengths, pinpoints areas for improvement, and generates time-optimized skill development plans—helping you 
                become interview-ready faster than traditional methods.
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # Feature Grid
        st.markdown("""
        <div class="feature-grid">
            <div class="feature-card">
                <div class="feature-icon">🎯</div>
                <h3 class="feature-title">ATS Optimization</h3>
                <p class="feature-desc">Advanced AI analysis to ensure your resume passes through Applicant Tracking Systems with higher scores.</p>
            </div>
            <div class="feature-card">
                <div class="feature-icon">📊</div>
                <h3 class="feature-title">Detailed Analytics</h3>
                <p class="feature-desc">Comprehensive breakdown of your resume's strengths, weaknesses, and compatibility scores.</p>
            </div>
            <div class="feature-card">
                <div class="feature-icon">🚀</div>
                <h3 class="feature-title">1-Week Learning</h3>
                <p class="feature-desc">Intensive learning paths to master missing skills in just one week with hands-on projects.</p>
            </div>
            <div class="feature-card">
                <div class="feature-icon">💼</div>
                <h3 class="feature-title">Project Generation</h3>
                <p class="feature-desc">Auto-generated projects to showcase both existing and new skills with impact metrics.</p>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)

else:
    # Logged in - Show main app
    st.markdown(f"""
    <div class="hero-section">
        <div class="hero-content">
            <h1 class="hero-title">AscendCV</h1>
            <p class="hero-subtitle">AI-Powered Resume Optimization with Personalized Skill Development</p>
            <div class="hero-features">
                <div class="feature-badge">🚀 Instant ATS Analysis</div>
                <div class="feature-badge">🎯 Personalized Learning</div>
                <div class="feature-badge">📊 Career Analytics</div>
                <div class="feature-badge">💼 Smart Recommendations</div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Sidebar with Dashboard and History
    with st.sidebar:
        st.markdown(f"""
        <div style="text-align: center; padding: 1.5rem; background: var(--gradient-primary); border-radius: 1.5rem; color: white; margin-bottom: 1.5rem; box-shadow: var(--shadow-lg);">
            <div style="font-size: 2.5rem; margin-bottom: 0.5rem;">👤</div>
            <h3>Welcome!</h3>
            <p style="font-size: 1.1rem; font-weight: 600;">{st.session_state.username}</p>
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("🚪 Logout", use_container_width=True):
            st.session_state.logged_in = False
            st.session_state.username = ""
            st.session_state.current_analysis = None
            st.rerun()
        
        st.markdown("---")
        
        # Dashboard Section in Sidebar
        st.markdown("### 📊 Dashboard")
        
        history = get_user_history(st.session_state.username)
        
        if history:
            df = pd.DataFrame(history, columns=['ID', 'Job Title', 'Match %', 'Date', 'Analysis Result'])
            df['Date'] = pd.to_datetime(df['Date']).dt.strftime('%Y-%m-%d')
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.metric("Total Analyses", len(history))
                avg_score = df['Match %'].mean()
                st.metric("Average Score", f"{avg_score:.1f}%")
            
            with col2:
                best_score = df['Match %'].max()
                st.metric("Best Score", f"{best_score:.1f}%")
                recent_count = len(df[df['Date'] >= (datetime.now() - timedelta(days=7)).strftime('%Y-%m-%d')])
                st.metric("This Week", recent_count)
            
            # Progress Timeline in Sidebar
            
        else:
            st.info("No analysis history yet")
        
        st.markdown("---")
        
        # History Section in Sidebar
        st.markdown("### 📋 History")
        
        if history:
            for idx, row in df.head(5).iterrows():
                with st.expander(f"{row['Job Title']} - {row['Match %']}%", expanded=False):
                    st.markdown(f"**Date:** {row['Date']}")
                    if st.button(f"Load Analysis", key=f"load_{row['ID']}", use_container_width=True):
                        analysis_data = get_analysis_by_id(row['ID'])
                        if analysis_data:
                            st.session_state.current_analysis = analysis_data
                            st.success("Analysis loaded! Switch to Resume Analysis tab.")
                            time.sleep(1)
                            st.rerun()
        else:
            st.info("No history yet")
        
        st.markdown("---")
        
        if st.session_state.current_analysis:
            st.markdown("### 🔍 Last Analysis")
            analysis = st.session_state.current_analysis
            col1, col2 = st.columns(2)
            with col1:
                st.metric("Match Score", f"{analysis['match_percentage']}%")
            with col2:
                st.metric("Missing Skills", len(analysis['missing_technical_skills']))

    # Main Content Tabs
    tab1, tab2, tab3, tab4 = st.tabs(["📊 Resume Analysis", "🎓 Skill Improvement", "💼 Auto Resume Generator", "📈 Progress Dashboard"])

    with tab1:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown("### 📝 Job Description Analysis")
        job_description = st.text_area(
            "Paste the job description:",
            height=200,
            placeholder="Copy and paste the full job description...",
            value=st.session_state.job_description
        )
        
        st.markdown("### 📄 Resume Upload")
        uploaded_file = st.file_uploader(
            "Upload your resume (PDF or DOCX):",
            type=['pdf', 'docx'],
        )
        
        job_title = st.text_input("🎯 Job Title (Optional):", placeholder="e.g., Senior Software Engineer")
        
        if st.button("🚀 Analyze Resume", type="primary", use_container_width=True):
            if not job_description.strip():
                st.warning("⚠️ Please provide a job description")
            elif not uploaded_file:
                st.warning("⚠️ Please upload your resume")
            else:
                with st.spinner("🔍 Performing AI analysis..."):
                    progress_bar = st.progress(0)
                    for i in range(100):
                        time.sleep(0.01)
                        progress_bar.progress(i + 1)
                    
                    resume_text = extract_text_from_file(uploaded_file)
                    
                    if not resume_text or len(resume_text.strip()) < 50:
                        st.error("❌ Could not extract text. Ensure file is text-based.")
                    else:
                        analyzer = st.session_state.analyzer
                        analysis = analyzer.analyze_resume_advanced(resume_text, job_description)
                        
                        save_analysis(
                            st.session_state.username,
                            job_title or "Unknown Position",
                            analysis['match_percentage'],
                            analysis
                        )
                        st.session_state.current_analysis = analysis
                        st.session_state.resume_text = resume_text
                        st.session_state.job_description = job_description
                        
                        st.success("✅ Analysis Complete!")
                        st.rerun()
        
        # Display results with enhanced profile summary
        if st.session_state.current_analysis:
            analysis = st.session_state.current_analysis
            
            st.markdown("## 📊 Analysis Results")
            
            # Enhanced Profile Summary Section
            st.markdown('<div class="profile-summary-card">', unsafe_allow_html=True)
            st.markdown("### 👤 Professional Profile Summary")
            
            # Calculate strength level
            score = analysis['match_percentage']
            if score >= 80:
                strength_level = "Excellent Match"
                strength_class = "strength-excellent"
            elif score >= 65:
                strength_level = "Good Match"
                strength_class = "strength-good"
            elif score >= 50:
                strength_level = "Moderate Match"
                strength_class = "strength-moderate"
            else:
                strength_level = "Skills Gap"
                strength_class = "strength-gap"
            
            # Strength level badge
            st.markdown(f'<div class="strength-badge {strength_class}" style="margin-bottom: 1.5rem;">{strength_level}</div>', unsafe_allow_html=True)
            
            # Calculate technical coverage
            tech_coverage = (analysis['technical_skills_matched'] / max(analysis['technical_skills_total'], 1)) * 100
            
            # Experience level assessment
            exp_years = analysis['experience_years']
            if exp_years >= 5:
                experience_level = "Senior Level"
            elif exp_years >= 3:
                experience_level = "Mid Level"
            elif exp_years >= 1:
                experience_level = "Junior Level"
            else:
                experience_level = "Entry Level"
            
            # Comprehensive paragraph summary
            profile_paragraph = f"""
            Your resume demonstrates a **{strength_level}** with the target position, 
            achieving an overall ATS compatibility score of **{analysis['match_percentage']}%**. With **{analysis['experience_years']} years of professional experience** 
            and expertise in **{analysis['technical_skills_matched']} out of {analysis['technical_skills_total']} required technical skills**, your profile shows 
            **{tech_coverage:.1f}% technical coverage**. 

            Your **{experience_level}** background is complemented by **{analysis['achievements_count']} quantified achievements** 
            that demonstrate tangible impact. The analysis indicates **{len(analysis['missing_technical_skills'])} key skill gaps** that present opportunities 
            for targeted development. 

            To enhance your competitiveness, focus on acquiring **{', '.join(analysis['missing_technical_skills'][:3]) if analysis['missing_technical_skills'] else 'additional relevant skills'}**. 
            Your current **keyword alignment score of {analysis['keyword_score']:.1f}%** suggests optimizing resume content with more job-specific terminology 
            could further improve ATS performance.
            """
            
            st.markdown(profile_paragraph)
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Visualizations Row
            col1, col2 = st.columns([1, 1])
            
            with col1:
                fig_gauge = create_advanced_match_gauge(analysis['match_percentage'])
                st.plotly_chart(fig_gauge, use_container_width=True)
            
            with col2:
                fig_radar = create_radar_chart(analysis)
                st.plotly_chart(fig_radar, use_container_width=True)
            
            # Score breakdown
            if 'score_breakdown' in analysis and analysis['score_breakdown']:
                st.markdown("### 📊 Score Breakdown Analysis")
                fig_breakdown = create_score_breakdown_chart(analysis['score_breakdown'])
                st.plotly_chart(fig_breakdown, use_container_width=True)
            
            # Key metrics
            st.markdown("### 📈 Key Performance Indicators")
            cols = st.columns(6)
            
            metrics = [
                ("Overall Match", f"{analysis['match_percentage']}%", "🎯"),
                ("Tech Skills", f"{analysis['technical_skills_matched']}/{analysis['technical_skills_total']}", "💻"),
                ("Soft Skills", f"{analysis['soft_skills_matched']}/{analysis['soft_skills_total']}", "🤝"),
                ("Experience", f"{analysis['experience_years']} yrs", "📅"),
                ("Keywords", f"{analysis['keyword_score']:.0f}%", "🔍"),
                ("Achievements", f"{analysis['achievements_count']}", "🏆")
            ]
            
            for col, (label, value, icon) in zip(cols, metrics):
                with col:
                    st.markdown(f"""
                    <div class="metric-card">
                        <div style="font-size: 1.5rem; margin-bottom: 0.3rem;">{icon}</div>
                        <span class="metric-value" style="font-size: 1.8rem;">{value}</span>
                        <div class="metric-label" style="font-size: 0.75rem;">{label}</div>
                    </div>
                    """, unsafe_allow_html=True)
            
            # Skill Match Chart
            st.markdown("### 🎯 Skill Match Comparison")
            fig_skill_match = create_skill_match_chart(
                analysis['technical_skills_matched'],
                analysis['technical_skills_total'],
                analysis['soft_skills_matched'],
                analysis['soft_skills_total']
            )
            st.plotly_chart(fig_skill_match, use_container_width=True)
            
            # Skills Comparison
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown('<div class="analysis-card">', unsafe_allow_html=True)
                st.markdown("#### ✅ Your Technical Skills")
                if analysis['resume_skills']['technical']:
                    for skill in analysis['resume_skills']['technical'][:15]:
                        st.markdown(f'<span class="skill-tag">{skill.title()}</span>', unsafe_allow_html=True)
                else:
                    st.info("No technical skills detected")
                st.markdown('</div>', unsafe_allow_html=True)
            
            with col2:
                st.markdown('<div class="analysis-card">', unsafe_allow_html=True)
                st.markdown("#### 🎯 Required Technical Skills")
                if analysis['job_skills']['technical']:
                    for skill in analysis['job_skills']['technical'][:15]:
                        st.markdown(f'<span class="skill-tag">{skill.title()}</span>', unsafe_allow_html=True)
                else:
                    st.info("No specific skills required")
                st.markdown('</div>', unsafe_allow_html=True)
            
            # Missing skills with chart
            if analysis['missing_technical_skills']:
                st.markdown('<div class="analysis-card">', unsafe_allow_html=True)
                st.markdown("### 🔍 Skill Gap Analysis")
                
                st.warning(f"**{len(analysis['missing_technical_skills'])} skills to develop**")
                st.info("💡 **Visit 'Skill Improvement' tab for personalized learning plans**")
                
                fig_missing = create_missing_skills_chart(analysis['missing_technical_skills'], analysis['improvement_plan'])
                if fig_missing:
                    st.plotly_chart(fig_missing, use_container_width=True)
                
                st.markdown("#### Missing Skills List:")
                for skill in analysis['missing_technical_skills'][:12]:
                    st.markdown(f'<span class="missing-skill-tag">{skill.title()}</span>', unsafe_allow_html=True)
                
                if len(analysis['missing_technical_skills']) > 12:
                    st.info(f"*... and {len(analysis['missing_technical_skills']) - 12} more*")
                
                st.markdown('</div>', unsafe_allow_html=True)
            else:
                st.success("🎉 Your skills align perfectly with the job requirements!")
            
            # Project Recommendations
            if analysis.get('generated_projects'):
                st.markdown('<div class="learning-path-card">', unsafe_allow_html=True)
                st.markdown("### 💡 Recommended Projects")
                
                for i, project in enumerate(analysis['generated_projects'], 1):
                    with st.expander(f"**{i}. {project['name']}** | Skills: {', '.join(project['skills'])}"):
                        st.write("**Project Description:**")
                        for bullet in project['bullets']:
                            st.write(f"• {bullet}")
                        
                        st.write("**Why this project?**")
                        st.write(f"This project helps you practice {len(project['skills'])} skills including both your existing strengths and areas for improvement.")
                
                st.markdown('</div>', unsafe_allow_html=True)
            
            # Achievements section
            if analysis.get('achievements'):
                st.markdown('<div class="analysis-card">', unsafe_allow_html=True)
                st.markdown("### 🏆 Quantified Achievements Found")
                for i, achievement in enumerate(analysis['achievements'][:5], 1):
                    st.markdown(f"{i}. {achievement}")
                st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    with tab2:
        st.markdown("## 🎓 One-Week Intensive Learning Path")
        st.markdown('<div class="card">', unsafe_allow_html=True)
        
        if st.session_state.current_analysis:
            analysis = st.session_state.current_analysis
            missing_skills = analysis['missing_technical_skills']
            improvement_plan = analysis['improvement_plan']
            
            if missing_skills:
                st.markdown("""
                ### 🚀 Accelerated Skill Development Plan
                
                **Complete these skills in ONE WEEK with our intensive learning schedule:**
                - **6-8 hours per day** dedicated learning
                - **Hands-on projects** from day 1
                - **Build portfolio projects** matching job requirements
                - **Practice with real-world scenarios**
                """)
                
                for skill, plan in improvement_plan.items():
                    st.markdown('<div class="learning-path-card">', unsafe_allow_html=True)
                    
                    with st.expander(f"📖 **{skill.title()}** - {plan.get('time_estimate', '1 week')} - Priority: {plan.get('priority', 'Medium')}"):
                        
                        st.markdown("#### 📅 7-Day Intensive Learning Schedule:")
                        for day in plan.get('learning_path', []):
                            st.write(f"• {day}")
                        
                        st.markdown("#### 💡 Project Ideas to Build:")
                        for project in plan.get('project_ideas', [])[:3]:
                            st.write(f"• {project}")
                        
                        if plan.get('resources'):
                            st.markdown("#### 📚 Learning Resources:")
                            for resource in plan.get('resources', [])[:3]:
                                st.markdown(f"• {resource}", unsafe_allow_html=True)
                        
                        if plan.get('practice_platforms'):
                            st.markdown("#### 🏆 Practice Platforms:")
                            for platform in plan.get('practice_platforms', [])[:3]:
                                st.markdown(f"• {platform}", unsafe_allow_html=True)
                        
                        if plan.get('quick_wins'):
                            st.markdown("#### ⚡ Quick Wins (First 2 Days):")
                            for tip in plan.get('quick_wins', [])[:3]:
                                st.write(f"• {tip}")
                    
                    st.markdown('</div>', unsafe_allow_html=True)
                
                # Weekly Schedule Overview
                st.markdown('<div class="quick-action-card">', unsafe_allow_html=True)
                st.subheader("📅 Weekly Study Plan Overview")
                
                days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
                for i, day in enumerate(days):
                    st.write(f"**{day}:** 6-8 hours focused learning + 2 hours practice")
                
                st.info("💡 **Pro Tip:** Build one complete project this weekend to solidify your learning!")
                st.markdown('</div>', unsafe_allow_html=True)
                
            else:
                st.success("🎉 Congratulations! Your skills already match the job requirements well.")
        else:
            st.info("👆 Complete a resume analysis first to see personalized skill development plans!")
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    with tab3:
        st.markdown("## 💼 Automated Resume Generator")
        st.markdown('<div class="card">', unsafe_allow_html=True)
        
        if st.session_state.current_analysis:
            analysis = st.session_state.current_analysis
            
            st.markdown("""
            ### 🚀 AI-Powered Resume Generation
            
            Our automated system will generate a perfectly optimized resume based on:
            - **Your uploaded resume content**
            - **Job description requirements** 
            - **ATS optimization techniques**
            - **Keyword matching analysis**
            - **Missing skills included in technical skills section**
            - **Project recommendations for skill demonstration**
            """)
            
            # Personal information override (optional)
            st.markdown("### 👤 Personal Information")
            col1, col2 = st.columns(2)
            with col1:
                auto_name = st.text_input("Full Name", value="", key="auto_name", placeholder="Your Name")
                auto_email = st.text_input("Email Address", value="", key="auto_email", placeholder="your.email@example.com")
            with col2:
                auto_phone = st.text_input("Phone Number", value="", key="auto_phone", placeholder="+1-234-567-8900")
                auto_location = st.text_input("Location", value="", key="auto_location", placeholder="City, State")
            
            # Resume customization options
            st.markdown("### ⚙️ Resume Customization")
            include_projects = st.checkbox("Include Project Suggestions", value=True,
                                        help="Add relevant project suggestions to demonstrate your skills")
            
            # Generate automated resume
            if st.button("💼 Generate Enhanced Resume", type="primary", use_container_width=True):
                with st.spinner("🔄 AI is generating your optimized resume..."):
                    # Prepare user info
                    user_info = {
                        'name': auto_name or "Your Name",
                        'email': auto_email or "your.email@example.com",
                        'phone': auto_phone or "+1-234-567-8900", 
                        'location': auto_location or "City, State"
                    }
                    
                    # Generate enhanced resume
                    enhanced_resume_buffer = st.session_state.resume_generator.generate_enhanced_resume(
                        analysis=analysis,
                        resume_text=st.session_state.resume_text,
                        job_description=st.session_state.job_description,
                        user_info=user_info,
                        include_projects=include_projects
                    )
                    
                    # Generate Word resume
                    word_resume_buffer = st.session_state.resume_generator.generate_word_resume(
                        analysis=analysis,
                        resume_text=st.session_state.resume_text,
                        job_description=st.session_state.job_description,
                        user_info=user_info,
                        include_projects=include_projects
                    )
                    
                    if enhanced_resume_buffer and word_resume_buffer:
                        st.success("✅ AI-generated resume created successfully!")
                        
                        # Show resume preview
                        st.markdown("### 👀 Enhanced Resume Preview")
                        
                        # Skills Overview
                        col1, col2 = st.columns(2)
                        with col1:
                            st.markdown("""
                            **✅ ENHANCED FEATURES**
                            - Technical skills matched with job requirements
                            - Missing skills included as 'Developing'  
                            - ATS-optimized keyword placement
                            - Real education information extracted
                            """)
                        
                        with col2:
                            if include_projects:
                                st.markdown("""
                                **💼 PROJECT HIGHLIGHTS**
                                - 2-3 relevant project suggestions
                                - Skills demonstration through projects
                                - Impact-focused project descriptions
                                """)
                        
                        # Download options
                        st.markdown("### 📥 Download Your Enhanced Resume")
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.download_button(
                                label="📄 Download PDF Resume",
                                data=enhanced_resume_buffer.getvalue(),
                                file_name=f"Enhanced_Resume_{user_info['name'].replace(' ', '_')}.pdf",
                                mime="application/pdf",
                                use_container_width=True
                            )
                        
                        with col2:
                            st.download_button(
                                label="📝 Download Word Resume", 
                                data=word_resume_buffer.getvalue(),
                                file_name=f"Enhanced_Resume_{user_info['name'].replace(' ', '_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        
                        # Show optimization details
                        st.markdown("### 🔍 Resume Optimization Details")
                        metrics_col1, metrics_col2, metrics_col3 = st.columns(3)
                        
                        with metrics_col1:
                            st.metric("Total Skills", f"{len(analysis['resume_skills']['technical']) + len(analysis['missing_technical_skills'])}")
                            st.metric("ATS Score", f"{analysis['match_percentage']}%")
                        
                        with metrics_col2:
                            st.metric("Projects Added", f"{min(3, len(analysis.get('generated_projects', [])))}")
                            st.metric("Skills Developing", len(analysis['missing_technical_skills']))
                        
                        with metrics_col3:
                            st.metric("Experience", f"{analysis['experience_years']} yrs")
                            st.metric("Education", "Extracted" if user_info.get('education') else "Default")
                            
                    else:
                        st.error("❌ Failed to generate resume. Please try again.")
        else:
            st.info("👆 Complete a resume analysis first to generate an enhanced resume!")
            
            st.markdown("""
            ### How It Works:
            1. **Upload your resume** and **job description** in the Resume Analysis tab
            2. **Get comprehensive analysis** of skills and compatibility  
            3. **Generate enhanced resume** automatically with AI
            4. **Download professionally formatted** resume in PDF or Word format
            
            ### The Enhanced Resume Includes:
            - **ATS-optimized formatting** and keywords
            - **Combined skills section** with existing and developing skills
            - **Real education information** extracted from your resume
            - **2-3 project suggestions** for skill demonstration
            - **Professional layout** that stands out to recruiters
            """)
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    with tab4:
        st.markdown("## 📈 Progress Dashboard")
        st.markdown('<div class="card">', unsafe_allow_html=True)
        
        history = get_user_history(st.session_state.username)
        
        if history:
            df = pd.DataFrame(history, columns=['ID', 'Job Title', 'Match %', 'Date', 'Analysis Result'])
            df['Date'] = pd.to_datetime(df['Date'])
            
            # Overall Statistics
            st.markdown("### 📊 Overall Performance")
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("Total Analyses", len(history))
            with col2:
                avg_score = df['Match %'].mean()
                st.metric("Average Score", f"{avg_score:.1f}%")
            with col3:
                best_score = df['Match %'].max()
                st.metric("Best Score", f"{best_score:.1f}%")
            with col4:
                recent_count = len(df[df['Date'] >= (datetime.now() - timedelta(days=7))])
                st.metric("This Week", recent_count)
            
            # Progress Timeline
            st.markdown("### 📈 Progress Timeline")
            fig_timeline = create_progress_timeline(history)
            if fig_timeline:
                st.plotly_chart(fig_timeline, use_container_width=True)
            
            # Recent Analyses Table
            st.markdown("### 📋 Recent Analyses")
            display_df = df.head(10).copy()
            display_df['Date'] = display_df['Date'].dt.strftime('%Y-%m-%d')
            st.dataframe(
                display_df[['Job Title', 'Match %', 'Date']],
                use_container_width=True,
                hide_index=True
            )
            
            # Improvement Tips
            if len(history) > 1:
                latest_score = df.iloc[0]['Match %']
                previous_score = df.iloc[1]['Match %']
                improvement = latest_score - previous_score
                
                st.markdown("### 💡 Improvement Insights")
                if improvement > 0:
                    st.success(f"🎉 Great progress! Your score improved by {improvement:.1f}% since your last analysis.")
                elif improvement < 0:
                    st.warning(f"📉 Your score decreased by {abs(improvement):.1f}%. Consider focusing on the skill gaps identified.")
                else:
                    st.info("📊 Your score remained the same. Try implementing the learning plan recommendations.")
        
        else:
            st.info("📊 No analysis history yet. Complete your first resume analysis to see your progress!")
            
            st.markdown("""
            ### What You'll See Here:
            - **Progress tracking** over time
            - **Performance metrics** and trends
            - **Improvement insights** and recommendations
            - **Skill development** progress
            - **Career growth** analytics
            """)
        
        st.markdown('</div>', unsafe_allow_html=True)


